"""
report_generator.py - Claude/LLM 분석 + Word 보고서 생성
OpenAI 호환 API 사용 (회사 내부 루코드 등)
"""
import os
from datetime import datetime
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import config


def _build_llm_client() -> OpenAI:
    """OpenAI 호환 클라이언트 생성"""
    config.validate_llm_config()
    return OpenAI(
        api_key=config.LLM_API_KEY,
        base_url=config.LLM_BASE_URL,
    )


def analyze_with_llm(content_text: str, title: str = "", context: str = "") -> str:
    """
    LLM으로 컨플루언스 내용 분석 및 보고서 초안 생성
    """
    client = _build_llm_client()

    system_prompt = """당신은 기업 업무 보고서 작성 전문가입니다.
주어진 Confluence 페이지 내용을 분석하여 경영진에게 보고할 수 있는
명확하고 간결한 보고서를 작성하세요.

보고서 형식:
1. 핵심 요약 (3-5문장)
2. 주요 내용 및 현황
3. 완료된 사항
4. 진행 중인 사항
5. 이슈 및 리스크
6. 다음 단계 / 액션 아이템

- 한국어로 작성
- 명확하고 간결하게
- 불릿 포인트 활용
- 구체적인 수치나 날짜 포함"""

    user_prompt = f"""다음 Confluence 페이지 내용을 분석하여 보고서를 작성해주세요.

제목: {title}
{f'추가 컨텍스트: {context}' if context else ''}

=== 내용 ===
{content_text[:6000]}
"""

    response = client.chat.completions.create(
        model=config.LLM_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
        max_tokens=config.LLM_MAX_TOKENS,
        temperature=0.3,
    )
    return response.choices[0].message.content


def analyze_multiple_pages(pages: list) -> str:
    """여러 페이지를 종합 분석"""
    client = _build_llm_client()

    combined = ""
    for row in pages:
        _, title, text, part, year, week, ptype = row
        combined += f"\n\n### {part} | {year} | {week} ({ptype})\n{text[:1500]}"

    system_prompt = """당신은 기업 업무 보고서 작성 전문가입니다.
여러 부서/파트의 Confluence 페이지 내용을 종합하여
전사 보고서를 작성하세요.

보고서 형식:
1. 전체 요약
2. 파트별 주요 현황
3. 공통 이슈 및 리스크
4. 종합 다음 단계
한국어로 명확하게 작성하세요."""

    response = client.chat.completions.create(
        model=config.LLM_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": f"다음 내용들을 종합 분석하여 보고서를 작성해주세요:\n{combined}"},
        ],
        max_tokens=config.LLM_MAX_TOKENS,
        temperature=0.3,
    )
    return response.choices[0].message.content


# ── Word 보고서 생성 ──────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int = 1):
    """헤딩 추가 (스타일 적용)"""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = "맑은 고딕"
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(16)
    else:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        run.font.size = Pt(13)
    return p


def _add_paragraph(doc: Document, text: str, bold: bool = False):
    """일반 단락 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(11)
    run.bold = bold
    return p


def create_word_report(
    analysis_text: str,
    title: str,
    output_path: str = None,
    meta: dict = None,
) -> str:
    """
    분석 결과를 Word 파일로 저장
    반환: 저장된 파일 경로
    """
    doc = Document()

    # 페이지 여백 설정 (A4)
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin  = section.bottom_margin = Cm(2.5)

    # 제목
    _add_heading(doc, title, level=1)

    # 메타 정보
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    meta_text = f"작성일시: {now}"
    if meta:
        if meta.get("part"):
            meta_text += f"  |  파트: {meta['part']}"
        if meta.get("year"):
            meta_text += f"  |  년도: {meta['year']}"
        if meta.get("week"):
            meta_text += f"  |  주차: {meta['week']}"
    p = doc.add_paragraph(meta_text)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_paragraph()  # 빈 줄

    # 분석 내용 파싱 및 삽입
    lines = analysis_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("## ") or line.startswith("# "):
            text = line.lstrip("#").strip()
            _add_heading(doc, text, level=2)
        elif line.startswith("### "):
            text = line.lstrip("#").strip()
            _add_heading(doc, text, level=3)
        elif line.startswith(("- ", "• ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(line[2:])
            run.font.name = "맑은 고딕"
            run.font.size = Pt(11)
        elif line.startswith(tuple(f"{i}." for i in range(1, 10))):
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(line.split(".", 1)[-1].strip())
            run.font.name = "맑은 고딕"
            run.font.size = Pt(11)
        else:
            _add_paragraph(doc, line)

    # 푸터
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = f"자동 생성 보고서  |  {now}"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # 저장
    if output_path is None:
        os.makedirs("reports", exist_ok=True)
        safe_title = "".join(c for c in title if c.isalnum() or c in " _-")[:30]
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        output_path = f"reports/{safe_title}_{ts}.docx"

    doc.save(output_path)
    print(f"✅ 보고서 저장 완료: {output_path}")
    return output_path
