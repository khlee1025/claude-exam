
"""
confluence_md_v2_ocr.py

Confluence child page collector + Markdown exporter + LLM report generator.

Key change from the original version:
  - Qwen is treated as a text model.
  - Images are downloaded, OCR'd, and the extracted text is interpreted by Qwen.
  - This mirrors the common Open WebUI flow where uploaded images are converted
    to text before being sent to a non-VL LLM.
"""

import base64
import glob
import io
import json
import os
import re
import subprocess
import sys
import threading
import traceback
from datetime import datetime
from typing import Dict, List, Optional
from urllib.parse import urljoin

import tkinter as tk
from tkinter import filedialog, messagebox, ttk


# ---------------------------------------------------------------------------
# Optional dependency loader
# ---------------------------------------------------------------------------

def _install(pkg: str):
    subprocess.check_call(
        [sys.executable, "-m", "pip", "install", pkg, "-q"],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )


def _import_or_install(import_name: str, pip_name: Optional[str] = None):
    try:
        return __import__(import_name)
    except ImportError:
        _install(pip_name or import_name)
        return __import__(import_name)


def _need_html2text():
    return _import_or_install("html2text")


def _need_beautifulsoup():
    bs4 = _import_or_install("bs4", "beautifulsoup4")
    return bs4.BeautifulSoup


def _need_openai_class():
    openai_mod = _import_or_install("openai")
    return openai_mod.OpenAI


# ---------------------------------------------------------------------------
# Palette
# ---------------------------------------------------------------------------

BLUE       = "#5A6C7D"
BLUE_DK    = "#4A5568"
BLUE_LT    = "#E2E8F0"
BG         = "#E8E9EB"
CARD       = "#F7F8FA"
BORDER     = "#D1D5DB"
TEXT       = "#2D3748"
TEXT_MUTED = "#718096"
LOG_BG     = "#2D3748"
LOG_FG     = "#E8E9EB"


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

BASE_URL      = os.getenv("CONFLUENCE_BASE_URL", "https://confluence.sec.samsung.net")
USER_DATA_DIR = os.getenv("CONFLUENCE_PROFILE_DIR", "./chrome_profile_confluence_md")
CONFIG_FILE   = os.path.join(os.path.dirname(os.path.abspath(__file__)), "llm_config.json")

# 기본값 초기화 (내부 서버 기본 연결)
LLM_API_KEY      = os.getenv("LLM_API_KEY", "")
LLM_BASE_URL     = os.getenv("LLM_BASE_URL", "http://10.240.246.158:8000/v1")
LLM_MODEL        = os.getenv("LLM_MODEL", "Qwen3.5-122B")
LLM_VISION_MODEL = os.getenv("LLM_VISION_MODEL", "Qwen3.5-122B")
LLM_MAX_TOKENS   = int(os.getenv("LLM_MAX_TOKENS", "4096"))




def load_llm_config():
    global LLM_API_KEY, LLM_BASE_URL, LLM_MODEL, LLM_VISION_MODEL, LLM_MAX_TOKENS
    if not os.path.exists(CONFIG_FILE):
        return
    try:
        with open(CONFIG_FILE, encoding="utf-8") as f:
            cfg = json.load(f)
        LLM_API_KEY = cfg.get("api_key", LLM_API_KEY)
        LLM_BASE_URL = cfg.get("base_url", LLM_BASE_URL)
        LLM_MODEL = cfg.get("model", LLM_MODEL)
        LLM_VISION_MODEL = cfg.get("vision_model", LLM_VISION_MODEL)
        LLM_MAX_TOKENS = int(cfg.get("max_tokens", LLM_MAX_TOKENS))
    except Exception as e:
        print(f"[설정 불러오기 실패] {e}")


def save_llm_config(api_key, base_url, model, vision_model, max_tokens):
    global LLM_API_KEY, LLM_BASE_URL, LLM_MODEL, LLM_VISION_MODEL, LLM_MAX_TOKENS
    LLM_API_KEY = api_key
    LLM_BASE_URL = base_url
    LLM_MODEL = model
    LLM_VISION_MODEL = vision_model
    LLM_MAX_TOKENS = int(max_tokens)
    with open(CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(
            {
                "api_key": api_key,
                "base_url": base_url,
                "model": model,
                "vision_model": vision_model,
                "max_tokens": int(max_tokens),
            },
            f,
            ensure_ascii=False,
            indent=2,
        )


load_llm_config()


def _llm():
    OpenAI = _need_openai_class()
    return OpenAI(
        api_key=LLM_API_KEY if LLM_API_KEY else "sk-ignored",
        base_url=LLM_BASE_URL,
    )


# ---------------------------------------------------------------------------
# OCR + LLM image analysis

def llm_analyze_image(image_bytes: bytes, context: str = "", mime_type: str = "image/png") -> str:
    """
    Vision API path:
      image bytes -> base64 -> Qwen Vision direct analysis
    """
    try:
        b64 = base64.b64encode(image_bytes).decode()
        data_url = f"data:{mime_type};base64,{b64}"
        resp = _llm().chat.completions.create(
            model=LLM_VISION_MODEL,
            temperature=0.2,
            max_tokens=min(1000, LLM_MAX_TOKENS),
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "당신은 기업 문서의 이미지, 표, 차트, 화면 캡처를 분석하는 전문가입니다.\n"
                                f"이미지 참고 정보: {context or '없음'}\n\n"
                                "이 이미지를 업무 보고서 관점에서 분석해주세요.\n\n"
                                "다음 항목을 최우선으로 추출하세요:\n"
                                "- 그래프/차트: 축 레이블, 수치, 단위, 추세\n"
                                "- 표: 행/열 데이터를 마크다운 표 형식으로 재현\n"
                                "- 수치/퍼센트/날짜: 정확히 그대로 기록\n"
                                "- 상태/색상 의미: 빨강=위험, 노랑=주의 등\n\n"
                                "읽기 어려운 글자는 '(판독불가)'로 표시하세요.\n"
                                "이미지가 단순 로고·아이콘이면 '단순 그래픽, 수치 없음'으로만 답하세요."
                            ),
                        },
                        {
                            "type": "image_url",
                            "image_url": {"url": data_url},
                        },
                    ],
                }
            ],
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"[Vision 이미지 분석 실패: {type(e).__name__}: {e}]"


def llm_summarize_page(title: str, markdown_text: str) -> str:
    try:
        resp = _llm().chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "당신은 기업 업무 보고서 작성 전문가입니다. "
                        "주어진 Confluence 페이지 내용을 분석하여 핵심만 요약하세요.\n"
                        "형식:\n"
                        "- **핵심 요약** (2~3 문장)\n"
                        "- **완료된 사항**\n"
                        "- **진행 중인 사항**\n"
                        "- **이슈 / 리스크**\n"
                        "한국어로 간결하게 작성. 없는 항목은 생략."
                    ),
                },
                {"role": "user", "content": f"페이지 제목: {title}\n\n내용:\n{markdown_text[:5000]}"},
            ],
            max_tokens=LLM_MAX_TOKENS,
            temperature=0.3,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"[LLM 요약 실패: {type(e).__name__}: {e}]"



# ──────────────────────────────────────────────────────────────────────────────
# 차트 자동 생성 (LLM 데이터 추출 → matplotlib PNG)
# ──────────────────────────────────────────────────────────────────────────────

def _import_matplotlib():
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.font_manager as fm
        # 한글 폰트 설정
        for fname in ["맑은 고딕", "NanumGothic", "AppleGothic", "DejaVu Sans"]:
            try:
                fm.findfont(fm.FontProperties(family=fname), fallback_to_default=False)
                matplotlib.rcParams["font.family"] = fname
                break
            except Exception:
                pass
        matplotlib.rcParams["axes.unicode_minus"] = False
        return plt
    except ImportError:
        return None


def llm_extract_chart_data(pages_content: list, callback=None) -> list:
    """MD 내용에서 실제 의미있는 수치 데이터를 추출 — 출처·해석·원본표 포함"""
    combined = "\n\n".join(
        f"### [{p['title']}]\n{p['content'][:3000]}" for p in pages_content
    )
    prompt = (
        "당신은 업무 보고서 분석 전문가입니다.\n"
        "아래 문서에서 실제로 의미있는 차트를 만들 수 있는 수치 데이터만 추출하세요.\n\n"
        "⚠️ 중요 규칙:\n"
        "- 문서에 명시적으로 존재하는 숫자/퍼센트/수량만 사용 (추측 금지)\n"
        "- 맥락 없이 나열된 숫자, 버전 번호, ID 등은 제외\n"
        "- 비교/추세/분포 의미가 명확한 데이터만 포함\n"
        "- 데이터가 없으면 반드시 빈 배열 [] 반환\n"
        "- 최대 4개까지만 추출\n\n"
        "JSON 배열만 반환 (마크다운 코드블록, 설명 없이):\n"
        "[\n"
        "  {\n"
        '    "type": "bar",       // bar(항목비교) | line(시계열추세) | pie(비율분포)\n'
        '    "title": "정확한 차트 제목 (문서 맥락 반영)",\n'
        '    "source": "출처 페이지/섹션명",\n'
        '    "labels": ["항목A", "항목B", "항목C"],\n'
        '    "values": [85, 60, 40],\n'
        '    "unit": "%",\n'
        '    "x_label": "X축 레이블 (해당하면)",\n'
        '    "y_label": "Y축 레이블 (해당하면)",\n'
        '    "insight": "이 데이터의 핵심 인사이트 한 문장 (예: Q3 매출이 전분기 대비 17% 증가)"\n'
        "  }\n"
        "]\n\n"
        "차트 타입 선택 기준:\n"
        "- bar: 여러 항목의 수치 비교 (진행률, 건수, 점수, 팀별 현황 등)\n"
        "- line: 시간 순서가 있는 추세 데이터 (월별, 분기별, 연도별)\n"
        "- pie: 전체 합이 의미있는 비율/구성 (상태별 분포, 비중)\n\n"
        f"=== 분석 대상 문서 ===\n{combined}"
    )
    try:
        if callback: callback("[차트] LLM 데이터 분석 중...")
        resp = _llm().chat.completions.create(
            model=LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2500,
            temperature=0.1,
        )
        raw = resp.choices[0].message.content.strip()
        m = re.search(r"\[.*?\]", raw, re.DOTALL)
        if m:
            charts = json.loads(m.group(0))
            valid  = [c for c in charts
                      if c.get("labels") and c.get("values")
                      and len(c["labels"]) == len(c["values"])
                      and len(c["values"]) >= 2
                      and all(isinstance(v, (int,float)) for v in c["values"])]
            if callback: callback(f"[차트] {len(valid)}개 유효 데이터셋 추출됨")
            return valid
    except Exception as e:
        if callback: callback(f"[차트] 데이터 추출 실패: {e}")
    return []


def generate_charts(chart_data: list, out_dir: str, callback=None) -> list:
    """chart_data 리스트 → PNG 저장. 최댓값 강조·변화율·추세선 포함."""
    plt = _import_matplotlib()
    if plt is None:
        if callback: callback("[차트] matplotlib 미설치. pip install matplotlib 필요.")
        return []

    import matplotlib.ticker as mticker
    try:
        import matplotlib.font_manager as fm
        # 한글 폰트 설정 (Windows/Mac/Linux 순)
        for fn in ["Malgun Gothic", "AppleGothic", "NanumGothic", "DejaVu Sans"]:
            try:
                plt.rcParams["font.family"] = fn
                plt.rcParams["axes.unicode_minus"] = False
                break
            except Exception:
                continue
    except Exception:
        pass

    os.makedirs(out_dir, exist_ok=True)
    BLUE    = "#1428A0"
    BLUE_LT = "#4B5FD6"
    BLUE_XL = "#E8EBF8"
    GRAY    = "#6B7280"
    RED     = "#DC2626"
    GREEN   = "#16A34A"
    COLORS  = [BLUE, BLUE_LT, "#7B8FE8", "#A8B4F0",
               "#2D9CDB", "#27AE60", "#F2994A", "#EB5757", "#9B51E0"]
    saved   = []

    for idx, spec in enumerate(chart_data):
        try:
            ctype   = spec.get("type", "bar")
            title   = spec.get("title", f"차트 {idx+1}")
            labels  = spec.get("labels", [])
            values  = spec.get("values", [])
            unit    = spec.get("unit", "")
            x_label = spec.get("x_label", "")
            y_label = spec.get("y_label", "")
            if not labels or not values or len(labels) != len(values):
                continue

            max_v = max(values); min_v = min(values); avg_v = sum(values)/len(values)
            max_i = values.index(max_v)

            fig, ax = plt.subplots(figsize=(9, 5))
            fig.patch.set_facecolor("white")

            # ── 파이 차트 ──────────────────────────────────────────────────
            if ctype == "pie":
                explode = [0.05 if i == max_i else 0 for i in range(len(values))]
                wedges, texts, autotexts = ax.pie(
                    values, labels=labels, autopct="%1.1f%%",
                    colors=COLORS[:len(values)], startangle=90,
                    pctdistance=0.80, explode=explode,
                    wedgeprops={"edgecolor": "white", "linewidth": 2})
                for at in autotexts:
                    at.set_fontsize(9); at.set_color("white"); at.set_fontweight("bold")
                ax.set_title(title, fontsize=13, fontweight="bold", color=BLUE, pad=16)

            # ── 라인 차트 ──────────────────────────────────────────────────
            elif ctype == "line":
                # 면 채우기
                ax.fill_between(range(len(labels)), values, alpha=0.08, color=BLUE)
                ax.plot(range(len(labels)), values, color=BLUE, linewidth=2.5,
                        marker="o", markersize=7, markerfacecolor="white",
                        markeredgecolor=BLUE, markeredgewidth=2, zorder=3)
                # 값 레이블
                for xi, (lbl, v) in enumerate(zip(labels, values)):
                    color = RED if xi == max_i else GRAY
                    ax.annotate(f"{v}{unit}", (xi, v),
                                textcoords="offset points", xytext=(0, 10),
                                ha="center", fontsize=9, color=color,
                                fontweight="bold" if xi == max_i else "normal")
                # 변화율 표시 (전 구간 대비)
                for xi in range(1, len(values)):
                    prev = values[xi-1]
                    if prev != 0:
                        chg = (values[xi] - prev) / abs(prev) * 100
                        sign = "▲" if chg > 0 else "▼"
                        col  = RED if chg > 0 else GREEN
                        ax.annotate(f"{sign}{abs(chg):.1f}%",
                                    xy=((xi-1+xi)/2, (values[xi-1]+values[xi])/2),
                                    fontsize=7.5, color=col, ha="center",
                                    fontweight="bold")
                # 평균선
                ax.axhline(avg_v, color=BLUE_LT, linewidth=1, linestyle="--", alpha=0.6)
                ax.text(len(labels)-0.5, avg_v, f" 평균 {avg_v:.1f}{unit}",
                        fontsize=8, color=BLUE_LT, va="bottom")
                ax.set_xticks(range(len(labels))); ax.set_xticklabels(labels, fontsize=10)
                ax.set_title(title, fontsize=13, fontweight="bold", color=BLUE)
                if x_label: ax.set_xlabel(x_label, fontsize=10, color=GRAY)
                if y_label: ax.set_ylabel(y_label, fontsize=10, color=GRAY)
                ax.set_facecolor("#F8F9FF")
                ax.grid(axis="y", linestyle="--", alpha=0.4, color="#C7CDE8")
                ax.spines[["top","right"]].set_visible(False)
                ax.tick_params(colors=GRAY)

            # ── 바 차트 (기본) ─────────────────────────────────────────────
            else:
                n = len(labels)
                use_horiz = n > 4  # 항목 많으면 가로 막대
                bar_colors = [RED if i == max_i else COLORS[i % len(COLORS)]
                              for i in range(n)]
                if use_horiz:
                    bars = ax.barh(labels, values, color=bar_colors,
                                   height=0.55, edgecolor="white", linewidth=0.8)
                    for bar, v in zip(bars, values):
                        ax.text(v + max_v*0.01, bar.get_y() + bar.get_height()/2,
                                f"{v}{unit}", va="center", fontsize=9.5,
                                color=RED if v == max_v else "#333",
                                fontweight="bold" if v == max_v else "normal")
                    ax.set_xlim(0, max_v * 1.18)
                    ax.set_facecolor("#F8F9FF")
                    ax.grid(axis="x", linestyle="--", alpha=0.4, color="#C7CDE8")
                    ax.spines[["top","right","left"]].set_visible(False)
                    ax.tick_params(left=False, labelsize=10)
                    if x_label: ax.set_xlabel(x_label, fontsize=10, color=GRAY)
                else:
                    bars = ax.bar(labels, values, color=bar_colors,
                                  width=0.55, edgecolor="white", linewidth=0.8)
                    for bar, v in zip(bars, values):
                        ax.text(bar.get_x() + bar.get_width()/2, v + max_v*0.01,
                                f"{v}{unit}", ha="center", fontsize=9.5,
                                color=RED if v == max_v else "#333",
                                fontweight="bold" if v == max_v else "normal")
                    ax.set_ylim(0, max_v * 1.18)
                    ax.set_facecolor("#F8F9FF")
                    ax.grid(axis="y", linestyle="--", alpha=0.4, color="#C7CDE8")
                    ax.spines[["top","right"]].set_visible(False)
                    ax.tick_params(bottom=False, labelsize=10)
                    if y_label: ax.set_ylabel(y_label, fontsize=10, color=GRAY)
                # 평균선
                ax.axhline(avg_v, color=BLUE_LT, linewidth=1.2, linestyle="--", alpha=0.7) if not use_horiz else                 ax.axvline(avg_v, color=BLUE_LT, linewidth=1.2, linestyle="--", alpha=0.7)
                ax.set_title(title, fontsize=13, fontweight="bold", color=BLUE, pad=12)

            plt.tight_layout(pad=1.8)
            fpath = os.path.join(out_dir, f"chart_{idx+1:02d}.png")
            plt.savefig(fpath, dpi=140, bbox_inches="tight",
                        facecolor="white", edgecolor="none")
            plt.close(fig)
            saved.append(fpath)
            if callback: callback(f"[차트] ✅ 생성: {title}")
        except Exception as e:
            if callback: callback(f"[차트] 생성 실패 [{idx+1}]: {e}")
            try: plt.close()
            except: pass

    return saved

def llm_generate_report(selected_md_contents: List[Dict]) -> str:
    combined = ""
    for item in selected_md_contents:
        combined += f"\n\n### {item['title']}\n{item['content']}"

    try:
        resp = _llm().chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "당신은 400 명 규모 조직의 팀장을 위한 종합 업무 보고서 작성 전문가입니다.\n\n"
                        "【작성 원칙】\n"
                        "1. 반드시 제공된 문서 내용만을 근거로 작성하세요. 추측하거나 내용을 창작하지 마세요.\n"
                        "2. 전문 용어가 나오면 괄호 안에 쉬운 말로 설명을 덧붙이세요.\n"
                        "3. 격식 있는 문어체로 작성하세요. (~하였습니다, ~진행 중에 있습니다)\n"
                        "4. 쉽고 구체적으로 풀어 쓰세요.\n"
                        "5. 수치, 날짜, 완료 여부 등 구체적인 사실은 그대로 포함하세요.\n\n"
                        "【보고서 구성】\n"
                        "# 1. 전체 요약\n"
                        "# 2. 항목별 상세 현황\n"
                        "# 3. 완료된 주요 사항\n"
                        "# 4. 진행 중인 주요 과제\n"
                        "# 5. 이슈 및 리스크\n"
                        "# 6. 다음 단계 / 액션 아이템\n\n"
                        "한국어 격식체로 작성. 문서에 없는 내용은 절대 추가하지 않음."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"아래 {len(selected_md_contents)}개 페이지의 내용을 바탕으로 종합 보고서를 작성해주세요.\n\n"
                        f"반드시 아래 제공된 내용만 사용하고, 없는 내용은 만들지 마세요.\n\n"
                        f"=== 각 페이지 내용 ===\n{combined}"
                    ),
                },
            ],
            max_tokens=LLM_MAX_TOKENS * 3,
            temperature=0.2,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"[보고서 생성 실패: {type(e).__name__}: {e}]"


# ---------------------------------------------------------------------------
# Markdown / Confluence helpers
# ---------------------------------------------------------------------------

def html_to_markdown(html: str, page_session=None, base_url: str = "", callback=None, img_save_dir: str = "") -> str:
    def _log(msg):
        if callback: callback(msg)
    _img_counter = [0]
    BeautifulSoup = _need_beautifulsoup()
    html2text = _need_html2text()
    soup = BeautifulSoup(html, "html.parser")

    imgs = soup.find_all("img")
    _log(f"[이미지 탐색] <img> 태그 발견: {len(imgs)}개")
    for img in imgs:
        src = img.get("src", "")
        alt = img.get("alt", "이미지")
        if not src:
            _log(f"  → src 없음, 스킵")
            continue

        img_desc = None
        if page_session and src:
            try:
                full_src = urljoin(base_url + "/", src) if not src.startswith("http") else src
                _log(f"  [이미지] src={src[:80]}")
                _log(f"  [이미지] full_url={full_src[:100]}")
                # page.request.get() 은 SSO/HTTPOnly 쿠키를 제대로 전달 못함
                # → 브라우저 컨텍스트의 fetch() 사용 (모든 쿠키 자동 포함)
                result = page_session.evaluate("""
                    async (url) => {
                        try {
                            const resp = await fetch(url, {credentials: 'include'});
                            if (!resp.ok) return {status: resp.status, data: null, contentType: ''};
                            const buf = await resp.arrayBuffer();
                            const bytes = Array.from(new Uint8Array(buf));
                            return {
                                status: resp.status,
                                data: bytes,
                                contentType: resp.headers.get('content-type') || ''
                            };
                        } catch(e) {
                            return {status: 0, data: null, contentType: '', error: e.message};
                        }
                    }
                """, full_src)

                status_code = result.get("status", 0)
                _log(f"  [이미지] HTTP {status_code}, contentType={result.get('contentType','')}, size={len(result.get('data') or [])} bytes")
                if result.get("status") == 200 and result.get("data"):
                    img_bytes = bytes(result["data"])
                    content_type = result.get("contentType", "image/png").split(";")[0].strip().lower()
                    if content_type == "image/jpg":
                        content_type = "image/jpeg"

                    if content_type in ("image/png", "image/jpeg", "image/webp", "image/gif") and len(img_bytes) > 300:
                        img_desc = llm_analyze_image(img_bytes, context=alt, mime_type=content_type)
                        # 이미지 파일 저장 (Word 삽입용)
                        if img_save_dir:
                            try:
                                os.makedirs(img_save_dir, exist_ok=True)
                                _img_counter[0] += 1
                                ext = content_type.split("/")[-1].replace("jpeg", "jpg")
                                img_filename = f"img_{_img_counter[0]:03d}.{ext}"
                                img_path = os.path.join(img_save_dir, img_filename)
                                with open(img_path, "wb") as _f:
                                    _f.write(img_bytes)
                                # MD에 상대 경로 마커 기록 (나중에 Word 삽입용)
                                rel_path = os.path.join("images", img_filename).replace("\\", "/")
                                img_desc = img_desc + f"\n\n[EMBED_IMAGE:{rel_path}]"
                                _log(f"  [이미지] 저장: {img_filename}")
                            except Exception as _e:
                                _log(f"  [이미지] 파일 저장 실패: {_e}")
                    else:
                        img_desc = f"[이미지 분석 스킵: content-type={content_type}, size={len(img_bytes)} bytes]"
                else:
                    status = result.get("status", 0)
                    err = result.get("error", "")
                    img_desc = f"[이미지 다운로드 실패: HTTP {status}{(' - ' + err) if err else ''}]"
            except Exception as e:
                img_desc = f"[이미지 다운로드 실패: {type(e).__name__}: {e}]"

        replacement = soup.new_tag("p")
        if img_desc:
            # [EMBED_IMAGE:path] 마커가 있으면 분리해서 별도 태그로
            if "[EMBED_IMAGE:" in img_desc:
                parts = img_desc.split("[EMBED_IMAGE:", 1)
                desc_text = parts[0].strip()
                embed_path = parts[1].rstrip("]").strip()
                replacement.string = f"**[이미지 분석: {alt}]** {desc_text}"
                img.replace_with(replacement)
                # EMBED_IMAGE 마커를 별도 p 태그로 삽입
                embed_tag = soup.new_tag("p")
                embed_tag.string = f"[EMBED_IMAGE:{embed_path}]"
                replacement.insert_after(embed_tag)
                continue
            else:
                replacement.string = f"**[이미지 분석: {alt}]** {img_desc}"
        else:
            replacement.string = f"**[이미지: {alt}]** (분석 불가 - src: {src[:120]})"
        img.replace_with(replacement)

    h = html2text.HTML2Text()
    h.ignore_links = False
    h.ignore_images = True
    h.body_width = 0
    h.protect_links = True
    h.wrap_links = False
    h.unicode_snob = True
    h.ignore_emphasis = False
    h.mark_code = True
    return h.handle(str(soup))


def clean_filename(title: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "_", title).strip()


def extract_page_id_from_url(url: str) -> Optional[str]:
    m = re.search(r"/pages/(\d+)", url)
    if m:
        return m.group(1)
    m = re.search(r"pageId=(\d+)", url)
    if m:
        return m.group(1)
    return None


def get_page_content(page_session, page_id: str) -> Optional[Dict]:
    try:
        resp = page_session.request.get(
            f"{BASE_URL}/rest/api/content/{page_id}",
            params={"expand": "body.view"},
        )
        if resp.status != 200:
            return None
        d = resp.json()
        return {
            "id": d.get("id"),
            "title": d.get("title", ""),
            "html": d.get("body", {}).get("view", {}).get("value", ""),
        }
    except Exception:
        return None


def get_children(page_session, parent_id: str) -> List[Dict]:
    pages, start = [], 0
    while True:
        resp = page_session.request.get(
            f"{BASE_URL}/rest/api/content/search",
            params={
                "cql": f"ancestor={parent_id} and type=page",
                "start": str(start),
                "limit": "50",
                "expand": "version",
            },
        )
        if resp.status != 200:
            break
        data = resp.json()
        results = data.get("results", [])
        if not results:
            break
        for doc in results:
            pages.append({"id": doc["id"], "title": doc["title"]})
        if len(pages) >= data.get("size", 0):
            break
        start += 50
        if start > 500:
            break
    return pages


def process_page(
    page_session,
    page_id: str,
    save_dir: str,
    depth: int,
    use_vision: bool,
    use_llm_summary: bool,
    callback=None,
):
    def log(msg):
        if callback:
            callback(msg)

    data = get_page_content(page_session, page_id)
    if not data:
        log(f"  [실패] 페이지 가져오기 실패: {page_id}")
        return

    title = data["title"]
    log(f"  처리 중: {title}")
    ps = page_session if use_vision else None
    img_dir = os.path.join(save_dir, "images") if use_vision else ""
    md_body = html_to_markdown(data["html"], page_session=ps, base_url=BASE_URL, callback=callback, img_save_dir=img_dir)

    llm_summary = ""
    if use_llm_summary:
        log(f"    LLM 요약 중: {title}")
        llm_summary = llm_summarize_page(title, md_body)

    md_lines = [
        f"# {title}",
        "",
        "---",
        f"페이지 ID: {page_id}",
        f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "---",
        "",
    ]
    if llm_summary:
        md_lines += ["## AI 요약", "", llm_summary, "", "---", ""]
    md_lines += ["## 원문 내용", "", md_body]

    safe = clean_filename(title)
    path = os.path.join(save_dir, f"{page_id}_{safe}.md")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))
    log(f"  [저장] {path}")

    if depth > 0:
        children = get_children(page_session, page_id)
        if children:
            child_dir = os.path.join(save_dir, f"sub_{safe}")
            os.makedirs(child_dir, exist_ok=True)
            for child in children:
                process_page(
                    page_session,
                    child["id"],
                    child_dir,
                    depth - 1,
                    use_vision,
                    use_llm_summary,
                    callback,
                )


# ──────────────────────────────────────────────────────────────────────────────
# Word 보고서 생성 (docx-js 기반, node.js 필요)
# ──────────────────────────────────────────────────────────────────────────────
import shutil as _shutil
import subprocess as _subprocess
import base64 as _base64

_MAKE_REPORT_JS_B64 = "LyoqCiAqIENvbmZsdWVuY2UgTUQg4oaSIFdvcmQg67O06rOg7IScIOyDneyEseq4sCAoU2Ftc3VuZyBTdHlsZSkKICogVXNhZ2U6CiAqICAgbm9kZSBtYWtlX3JlcG9ydC5qcyA8aW5wdXQubWR8Zm9sZGVyPiA8b3V0cHV0LmRvY3g+IFt0aXRsZV0KICovCiJ1c2Ugc3RyaWN0IjsKY29uc3QgZnMgICA9IHJlcXVpcmUoImZzIik7CmNvbnN0IHBhdGggPSByZXF1aXJlKCJwYXRoIik7CmNvbnN0IHsKICBEb2N1bWVudCwgUGFja2VyLCBQYXJhZ3JhcGgsIFRleHRSdW4sIFRhYmxlLCBUYWJsZVJvdywgVGFibGVDZWxsLAogIEhlYWRlciwgRm9vdGVyLCBBbGlnbm1lbnRUeXBlLCBIZWFkaW5nTGV2ZWwsIEJvcmRlclN0eWxlLCBXaWR0aFR5cGUsCiAgU2hhZGluZ1R5cGUsIFZlcnRpY2FsQWxpZ24sIFBhZ2VOdW1iZXIsIFBhZ2VCcmVhaywgVGFibGVPZkNvbnRlbnRzLAogIExldmVsRm9ybWF0LCBJbWFnZVJ1biwKfSA9IHJlcXVpcmUoImRvY3giKTsKCi8vIOKUgOKUgOKUgCDsg4nsg4Eg4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSACmNvbnN0IEMgPSB7CiAgYmx1ZTogICAgICAiMTQyOEEwIiwKICBsaWdodEJsdWU6ICJFOEVCRjgiLAogIG1pZEJsdWU6ICAgIkM3Q0RFOCIsCiAgcm93QWx0OiAgICAiRjRGNkZEIiwKICBncmF5OiAgICAgICI2QjcyODAiLAogIHJlZDogICAgICAgIkRDMjYyNiIsCiAgZ3JlZW46ICAgICAiMTZBMzRBIiwKICB3aGl0ZTogICAgICJGRkZGRkYiLAogIGRhcms6ICAgICAgIjFGMjkzNyIsCn07CgovLyDilIDilIDilIAg7Jyg7Yu4IOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgApmdW5jdGlvbiBib3JkZXIoY29sb3IgPSBDLm1pZEJsdWUsIHNpemUgPSA0KSB7CiAgY29uc3QgYiA9IHsgc3R5bGU6IEJvcmRlclN0eWxlLlNJTkdMRSwgc2l6ZSwgY29sb3IgfTsKICByZXR1cm4geyB0b3A6IGIsIGJvdHRvbTogYiwgbGVmdDogYiwgcmlnaHQ6IGIgfTsKfQoKZnVuY3Rpb24gcGFyc2VJbmxpbmUodGV4dCwgc2l6ZSA9IDIyLCBkZWZhdWx0Q29sb3IgPSBDLmRhcmspIHsKICBjb25zdCBydW5zID0gW107CiAgY29uc3QgcmUgICA9IC8oXCpcKiguKz8pXCpcKnxcKiguKz8pXCp8YChbXmBdKylgKS9nOwogIGxldCBsYXN0ID0gMCwgbTsKICB3aGlsZSAoKG0gPSByZS5leGVjKHRleHQpKSAhPT0gbnVsbCkgewogICAgaWYgKG0uaW5kZXggPiBsYXN0KQogICAgICBydW5zLnB1c2gobmV3IFRleHRSdW4oeyB0ZXh0OiB0ZXh0LnNsaWNlKGxhc3QsIG0uaW5kZXgpLCBzaXplLCBmb250OiAi66eR7J2AIOqzoOuUlSIsIGNvbG9yOiBkZWZhdWx0Q29sb3IgfSkpOwogICAgaWYgICAgICAobVsyXSkgcnVucy5wdXNoKG5ldyBUZXh0UnVuKHsgdGV4dDogbVsyXSwgYm9sZDogdHJ1ZSwgICAgc2l6ZSwgZm9udDogIuunkeydgCDqs6DrlJUiLCBjb2xvcjogZGVmYXVsdENvbG9yIH0pKTsKICAgIGVsc2UgaWYgKG1bM10pIHJ1bnMucHVzaChuZXcgVGV4dFJ1bih7IHRleHQ6IG1bM10sIGl0YWxpY3M6IHRydWUsIHNpemUsIGZvbnQ6ICLrp5HsnYAg6rOg65SVIiwgY29sb3I6IGRlZmF1bHRDb2xvciB9KSk7CiAgICBlbHNlIGlmIChtWzRdKSBydW5zLnB1c2gobmV3IFRleHRSdW4oeyB0ZXh0OiBtWzRdLCBmb250OiAiQ29uc29sYXMiLCBzaXplOiBzaXplIC0gMiwKICAgICAgc2hhZGluZzogeyB0eXBlOiBTaGFkaW5nVHlwZS5DTEVBUiwgZmlsbDogIkYzRjRGNiIgfSwgY29sb3I6ICJCOTFDMUMiIH0pKTsKICAgIGxhc3QgPSBtLmluZGV4ICsgbVswXS5sZW5ndGg7CiAgfQogIGlmIChsYXN0IDwgdGV4dC5sZW5ndGgpCiAgICBydW5zLnB1c2gobmV3IFRleHRSdW4oeyB0ZXh0OiB0ZXh0LnNsaWNlKGxhc3QpLCBzaXplLCBmb250OiAi66eR7J2AIOqzoOuUlSIsIGNvbG9yOiBkZWZhdWx0Q29sb3IgfSkpOwogIHJldHVybiBydW5zLmxlbmd0aCA/IHJ1bnMgOiBbbmV3IFRleHRSdW4oeyB0ZXh0LCBzaXplLCBmb250OiAi66eR7J2AIOqzoOuUlSIsIGNvbG9yOiBkZWZhdWx0Q29sb3IgfSldOwp9CgovLyDilIDilIDilIAg7Zek65SpIOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgApjb25zdCBIRUFESU5HX0NGRyA9IHsKICAxOiB7IGxldmVsOiBIZWFkaW5nTGV2ZWwuSEVBRElOR18xLCBzaXplOiAzNiwgY29sb3I6IEMuYmx1ZSwgIGJlZm9yZTogNDAwLCBhZnRlcjogMjAwLCBvdXRsaW5lOiAwIH0sCiAgMjogeyBsZXZlbDogSGVhZGluZ0xldmVsLkhFQURJTkdfMiwgc2l6ZTogMzAsIGNvbG9yOiBDLmJsdWUsICBiZWZvcmU6IDMwMCwgYWZ0ZXI6IDE2MCwgb3V0bGluZTogMSB9LAogIDM6IHsgbGV2ZWw6IEhlYWRpbmdMZXZlbC5IRUFESU5HXzMsIHNpemU6IDI2LCBjb2xvcjogIjMzMzMzMyIsYmVmb3JlOiAyNDAsIGFmdGVyOiAxMjAsIG91dGxpbmU6IDIgfSwKICA0OiB7IGxldmVsOiBIZWFkaW5nTGV2ZWwuSEVBRElOR180LCBzaXplOiAyNCwgY29sb3I6ICI1NTU1NTUiLGJlZm9yZTogMTgwLCBhZnRlcjogIDgwLCBvdXRsaW5lOiAzIH0sCn07CmZ1bmN0aW9uIG1ha2VIZWFkaW5nKHRleHQsIGRlcHRoKSB7CiAgY29uc3QgYyA9IEhFQURJTkdfQ0ZHW01hdGgubWluKGRlcHRoLCA0KV0gfHwgSEVBRElOR19DRkdbNF07CiAgcmV0dXJuIG5ldyBQYXJhZ3JhcGgoewogICAgaGVhZGluZzogYy5sZXZlbCwKICAgIHNwYWNpbmc6IHsgYmVmb3JlOiBjLmJlZm9yZSwgYWZ0ZXI6IGMuYWZ0ZXIgfSwKICAgIGNoaWxkcmVuOiBbbmV3IFRleHRSdW4oeyB0ZXh0LCBib2xkOiB0cnVlLCBzaXplOiBjLnNpemUsIGNvbG9yOiBjLmNvbG9yLCBmb250OiAi66eR7J2AIOqzoOuUlSIgfSldLAogIH0pOwp9CgovLyDilIDilIDilIAg7ZGcIOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgApmdW5jdGlvbiBtYWtlVGFibGUoaGVhZGVycywgcm93cykgewogIGNvbnN0IGNvbHMgICA9IE1hdGgubWF4KGhlYWRlcnMubGVuZ3RoLCAuLi5yb3dzLm1hcChyID0+IHIubGVuZ3RoKSwgMSk7CiAgY29uc3QgdG90YWxXID0gOTAyNjsKICBjb25zdCBjb2xXICAgPSBNYXRoLmZsb29yKHRvdGFsVyAvIGNvbHMpOwogIGNvbnN0IGNvbFdzICA9IEFycmF5KGNvbHMpLmZpbGwoY29sVyk7CgogIGNvbnN0IGhkclJvdyA9IG5ldyBUYWJsZVJvdyh7CiAgICB0YWJsZUhlYWRlcjogdHJ1ZSwKICAgIGNoaWxkcmVuOiBoZWFkZXJzLm1hcCgoaCwgaSkgPT4gbmV3IFRhYmxlQ2VsbCh7CiAgICAgIHdpZHRoOiAgIHsgc2l6ZTogY29sV3NbaV0sIHR5cGU6IFdpZHRoVHlwZS5EWEEgfSwKICAgICAgYm9yZGVyczogYm9yZGVyKEMuYmx1ZSksCiAgICAgIHNoYWRpbmc6IHsgdHlwZTogU2hhZGluZ1R5cGUuQ0xFQVIsIGZpbGw6IEMuYmx1ZSB9LAogICAgICBtYXJnaW5zOiB7IHRvcDogMTAwLCBib3R0b206IDEwMCwgbGVmdDogMTYwLCByaWdodDogMTYwIH0sCiAgICAgIHZlcnRpY2FsQWxpZ246IFZlcnRpY2FsQWxpZ24uQ0VOVEVSLAogICAgICBjaGlsZHJlbjogW25ldyBQYXJhZ3JhcGgoeyBhbGlnbm1lbnQ6IEFsaWdubWVudFR5cGUuQ0VOVEVSLAogICAgICAgIGNoaWxkcmVuOiBbbmV3IFRleHRSdW4oeyB0ZXh0OiBoLCBib2xkOiB0cnVlLCBjb2xvcjogQy53aGl0ZSwgc2l6ZTogMjAsIGZvbnQ6ICLrp5HsnYAg6rOg65SVIiB9KV0gfSldLAogICAgfSkpLAogIH0pOwoKICBjb25zdCBkYXRhUm93cyA9IHJvd3MubWFwKChyb3csIHJpKSA9PgogICAgbmV3IFRhYmxlUm93KHsKICAgICAgY2hpbGRyZW46IEFycmF5KGNvbHMpLmZpbGwobnVsbCkubWFwKChfLCBjaSkgPT4gewogICAgICAgIGNvbnN0IHZhbCAgID0gKHJvd1tjaV0gfHwgIiIpLnRyaW0oKTsKICAgICAgICBjb25zdCBpc1JlZCA9IC/sp4Dsl7B87LSI6rO8fOyLpO2MqHzsmKTrpZh86rK96rOgLy50ZXN0KHZhbCk7CiAgICAgICAgY29uc3QgaXNHcm4gPSAv7KCV7IOBfOyZhOujjHzshLHqs7UvLnRlc3QodmFsKTsKICAgICAgICByZXR1cm4gbmV3IFRhYmxlQ2VsbCh7CiAgICAgICAgICB3aWR0aDogICB7IHNpemU6IGNvbFdzW2NpXSwgdHlwZTogV2lkdGhUeXBlLkRYQSB9LAogICAgICAgICAgYm9yZGVyczogYm9yZGVyKEMubWlkQmx1ZSksCiAgICAgICAgICBzaGFkaW5nOiB7IHR5cGU6IFNoYWRpbmdUeXBlLkNMRUFSLCBmaWxsOiByaSAlIDIgPT09IDAgPyBDLndoaXRlIDogQy5yb3dBbHQgfSwKICAgICAgICAgIG1hcmdpbnM6IHsgdG9wOiA4MCwgYm90dG9tOiA4MCwgbGVmdDogMTYwLCByaWdodDogMTYwIH0sCiAgICAgICAgICB2ZXJ0aWNhbEFsaWduOiBWZXJ0aWNhbEFsaWduLkNFTlRFUiwKICAgICAgICAgIGNoaWxkcmVuOiBbbmV3IFBhcmFncmFwaCh7IGNoaWxkcmVuOiBbbmV3IFRleHRSdW4oewogICAgICAgICAgICB0ZXh0OiB2YWwsIHNpemU6IDIwLCBmb250OiAi66eR7J2AIOqzoOuUlSIsCiAgICAgICAgICAgIGNvbG9yOiBpc1JlZCA/IEMucmVkIDogaXNHcm4gPyBDLmdyZWVuIDogQy5kYXJrLAogICAgICAgICAgICBib2xkOiBpc1JlZCB8fCBpc0dybiwKICAgICAgICAgIH0pXSB9KV0sCiAgICAgICAgfSk7CiAgICAgIH0pLAogICAgfSkKICApOwoKICByZXR1cm4gbmV3IFRhYmxlKHsgd2lkdGg6IHsgc2l6ZTogdG90YWxXLCB0eXBlOiBXaWR0aFR5cGUuRFhBIH0sIGNvbHVtbldpZHRoczogY29sV3MsIHJvd3M6IFtoZHJSb3csIC4uLmRhdGFSb3dzXSB9KTsKfQoKLy8g4pSA4pSA4pSAIOy9nOyVhOybgyAo7J2066+47KeAIOu2hOyEnSkg4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSACmZ1bmN0aW9uIG1ha2VDYWxsb3V0KHRleHQpIHsKICBjb25zdCBjbGVhbiA9IHRleHQucmVwbGFjZSgvXCpcKlxb7J2066+47KeAW15cXV0qXF1cKlwqXHMqL2csICIiKS5yZXBsYWNlKC9cKlwqL2csICIiKS50cmltKCk7CiAgcmV0dXJuIG5ldyBQYXJhZ3JhcGgoewogICAgc3BhY2luZzogeyBiZWZvcmU6IDEyMCwgYWZ0ZXI6IDEyMCB9LAogICAgaW5kZW50OiAgeyBsZWZ0OiAzNjAgfSwKICAgIGJvcmRlcjogIHsgbGVmdDogeyBzdHlsZTogQm9yZGVyU3R5bGUuVEhJQ0ssIHNpemU6IDIwLCBjb2xvcjogQy5ibHVlLCBzcGFjZTogOCB9IH0sCiAgICBzaGFkaW5nOiB7IHR5cGU6IFNoYWRpbmdUeXBlLkNMRUFSLCBmaWxsOiBDLmxpZ2h0Qmx1ZSB9LAogICAgY2hpbGRyZW46IFsKICAgICAgbmV3IFRleHRSdW4oeyB0ZXh0OiAi8J+UjSDsnbTrr7jsp4Ag67aE7ISdICAiLCBib2xkOiB0cnVlLCBzaXplOiAyMCwgZm9udDogIuunkeydgCDqs6DrlJUiLCBjb2xvcjogQy5ibHVlIH0pLAogICAgICBuZXcgVGV4dFJ1bih7IHRleHQ6IGNsZWFuLCBzaXplOiAyMCwgZm9udDogIuunkeydgCDqs6DrlJUiLCBjb2xvcjogIjJEM0E4QSIsIGl0YWxpY3M6IHRydWUgfSksCiAgICBdLAogIH0pOwp9CgovLyDilIDilIDilIAg7Y6Y7J207KeAIOy2nOyymCDrsLDrhIgg4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSACmZ1bmN0aW9uIG1ha2VTb3VyY2VCYW5uZXIocGFnZVRpdGxlKSB7CiAgcmV0dXJuIG5ldyBQYXJhZ3JhcGgoewogICAgc3BhY2luZzogeyBiZWZvcmU6IDAsIGFmdGVyOiAxNjAgfSwKICAgIHNoYWRpbmc6IHsgdHlwZTogU2hhZGluZ1R5cGUuQ0xFQVIsIGZpbGw6IEMubGlnaHRCbHVlIH0sCiAgICBib3JkZXI6ICB7IGxlZnQ6IHsgc3R5bGU6IEJvcmRlclN0eWxlLlRISUNLLCBzaXplOiAxNiwgY29sb3I6IEMuYmx1ZSwgc3BhY2U6IDYgfSB9LAogICAgaW5kZW50OiAgeyBsZWZ0OiAxNjAgfSwKICAgIGNoaWxkcmVuOiBbbmV3IFRleHRSdW4oeyB0ZXh0OiBg8J+ThCAgJHtwYWdlVGl0bGV9YCwgc2l6ZTogMTksIGZvbnQ6ICLrp5HsnYAg6rOg65SVIiwgY29sb3I6IEMuZ3JheSwgaXRhbGljczogdHJ1ZSB9KV0sCiAgfSk7Cn0KCi8vIOKUgOKUgOKUgCBNRCDihpIgY2hpbGRyZW4g67OA7ZmYIOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgApmdW5jdGlvbiBwYXJzZU1EKG1kLCBfc291cmNlUGF0aCkgewogIGNvbnN0IGNoaWxkcmVuID0gW107CiAgY29uc3QgbGluZXMgICAgPSBtZC5zcGxpdCgiXG4iKTsKICBsZXQgaSA9IDAsIGluQ29kZSA9IGZhbHNlLCBjb2RlTGluZXMgPSBbXTsKICBsZXQgdGhkcnMgPSBudWxsLCB0cm93cyA9IFtdOwoKICBmdW5jdGlvbiBmbHVzaFRhYmxlKCkgewogICAgaWYgKCF0aGRycykgcmV0dXJuOwogICAgY2hpbGRyZW4ucHVzaChtYWtlVGFibGUodGhkcnMsIHRyb3dzKSk7CiAgICBjaGlsZHJlbi5wdXNoKG5ldyBQYXJhZ3JhcGgoeyBzcGFjaW5nOiB7IGFmdGVyOiAxMjAgfSwgY2hpbGRyZW46IFtdIH0pKTsKICAgIHRoZHJzID0gbnVsbDsgdHJvd3MgPSBbXTsKICB9CgogIHdoaWxlIChpIDwgbGluZXMubGVuZ3RoKSB7CiAgICBjb25zdCBsaW5lID0gbGluZXNbaV0sIHMgPSBsaW5lLnRyaW0oKTsKCiAgICAvLyDsvZTrk5wg67iU66GdCiAgICBpZiAocy5zdGFydHNXaXRoKCJgYGAiKSkgewogICAgICBpZiAoIWluQ29kZSkgeyBpbkNvZGUgPSB0cnVlOyBjb2RlTGluZXMgPSBbXTsgaSsrOyBjb250aW51ZTsgfQogICAgICBpbkNvZGUgPSBmYWxzZTsKICAgICAgY2hpbGRyZW4ucHVzaChuZXcgUGFyYWdyYXBoKHsKICAgICAgICBzcGFjaW5nOiB7IGJlZm9yZTogODAsIGFmdGVyOiA4MCB9LAogICAgICAgIHNoYWRpbmc6IHsgdHlwZTogU2hhZGluZ1R5cGUuQ0xFQVIsIGZpbGw6ICJGM0Y0RjYiIH0sCiAgICAgICAgYm9yZGVyOiAgYm9yZGVyKEMubWlkQmx1ZSwgMiksCiAgICAgICAgaW5kZW50OiAgeyBsZWZ0OiAyNDAgfSwKICAgICAgICBjaGlsZHJlbjogW25ldyBUZXh0UnVuKHsgdGV4dDogY29kZUxpbmVzLmpvaW4oIlxuIiksIGZvbnQ6ICJDb25zb2xhcyIsIHNpemU6IDE4LCBjb2xvcjogIjFGMjkzNyIgfSldLAogICAgICB9KSk7CiAgICAgIGkrKzsgY29udGludWU7CiAgICB9CiAgICBpZiAoaW5Db2RlKSB7IGNvZGVMaW5lcy5wdXNoKGxpbmUpOyBpKys7IGNvbnRpbnVlOyB9CgogICAgLy8g7Zek65SpCiAgICBjb25zdCBobSA9IHMubWF0Y2goL14oI3sxLDR9KVxzKyguKykvKTsKICAgIGlmIChobSkgeyBmbHVzaFRhYmxlKCk7IGNoaWxkcmVuLnB1c2gobWFrZUhlYWRpbmcoaG1bMl0sIGhtWzFdLmxlbmd0aCkpOyBpKys7IGNvbnRpbnVlOyB9CgogICAgLy8g7ZGcCiAgICBpZiAocy5zdGFydHNXaXRoKCJ8IikpIHsKICAgICAgY29uc3QgY2VsbHMgPSBzLnNwbGl0KCJ8Iikuc2xpY2UoMSwgLTEpLm1hcChjID0+IGMudHJpbSgpKTsKICAgICAgaWYgKGNlbGxzLmV2ZXJ5KGMgPT4gL15bLTogXSskLy50ZXN0KGMpKSkgeyBpKys7IGNvbnRpbnVlOyB9CiAgICAgIGlmICghdGhkcnMpIHRoZHJzID0gY2VsbHM7IGVsc2UgdHJvd3MucHVzaChjZWxscyk7CiAgICAgIGkrKzsgY29udGludWU7CiAgICB9IGVsc2UgeyBmbHVzaFRhYmxlKCk7IH0KCiAgICAvLyDsnbTrr7jsp4Ag67aE7ISdIOy9nOyVhOybgwogICAgaWYgKHMuc3RhcnRzV2l0aCgiKipb7J2066+47KeAIikgfHwgcy5zdGFydHNXaXRoKCIqKltWaXNpb24iKSkgewogICAgICBjaGlsZHJlbi5wdXNoKG1ha2VDYWxsb3V0KHMpKTsgaSsrOyBjb250aW51ZTsKICAgIH0KCiAgICAvLyDquIDrqLjrpqwKICAgIGNvbnN0IGJtID0gcy5tYXRjaCgvXlstKitdXHMrKC4rKS8pOwogICAgaWYgKGJtKSB7CiAgICAgIGNoaWxkcmVuLnB1c2gobmV3IFBhcmFncmFwaCh7CiAgICAgICAgc3BhY2luZzogeyBiZWZvcmU6IDQwLCBhZnRlcjogNDAgfSwKICAgICAgICBpbmRlbnQ6ICB7IGxlZnQ6IDQ4MCwgaGFuZ2luZzogMjQwIH0sCiAgICAgICAgY2hpbGRyZW46IFtuZXcgVGV4dFJ1bih7IHRleHQ6ICLigKIgICIsIHNpemU6IDIyLCBmb250OiAi66eR7J2AIOqzoOuUlSIsIGNvbG9yOiBDLmJsdWUsIGJvbGQ6IHRydWUgfSksIC4uLnBhcnNlSW5saW5lKGJtWzFdKV0sCiAgICAgIH0pKTsKICAgICAgaSsrOyBjb250aW51ZTsKICAgIH0KCiAgICAvLyDrsojtmLgg66qp66GdCiAgICBjb25zdCBubSA9IHMubWF0Y2goL14oXGQrKVwuXHMrKC4rKS8pOwogICAgaWYgKG5tKSB7CiAgICAgIGNoaWxkcmVuLnB1c2gobmV3IFBhcmFncmFwaCh7CiAgICAgICAgc3BhY2luZzogeyBiZWZvcmU6IDQwLCBhZnRlcjogNDAgfSwKICAgICAgICBpbmRlbnQ6ICB7IGxlZnQ6IDQ4MCwgaGFuZ2luZzogMjgwIH0sCiAgICAgICAgY2hpbGRyZW46IFtuZXcgVGV4dFJ1bih7IHRleHQ6IGAke25tWzFdfS4gIGAsIHNpemU6IDIyLCBmb250OiAi66eR7J2AIOqzoOuUlSIsIGNvbG9yOiBDLmJsdWUsIGJvbGQ6IHRydWUgfSksIC4uLnBhcnNlSW5saW5lKG5tWzJdKV0sCiAgICAgIH0pKTsKICAgICAgaSsrOyBjb250aW51ZTsKICAgIH0KCiAgICAvLyDsnbjsmqkKICAgIGNvbnN0IHFtID0gcy5tYXRjaCgvXj5ccyooLiopLyk7CiAgICBpZiAocW0pIHsKICAgICAgY2hpbGRyZW4ucHVzaChuZXcgUGFyYWdyYXBoKHsKICAgICAgICBzcGFjaW5nOiB7IGJlZm9yZTogODAsIGFmdGVyOiA4MCB9LCBpbmRlbnQ6IHsgbGVmdDogNDgwIH0sCiAgICAgICAgYm9yZGVyOiAgeyBsZWZ0OiB7IHN0eWxlOiBCb3JkZXJTdHlsZS5TSU5HTEUsIHNpemU6IDE2LCBjb2xvcjogQy5ibHVlLCBzcGFjZTogOCB9IH0sCiAgICAgICAgc2hhZGluZzogeyB0eXBlOiBTaGFkaW5nVHlwZS5DTEVBUiwgZmlsbDogQy5saWdodEJsdWUgfSwKICAgICAgICBjaGlsZHJlbjogcGFyc2VJbmxpbmUocW1bMV0sIDIwKSwKICAgICAgfSkpOwogICAgICBpKys7IGNvbnRpbnVlOwogICAgfQoKICAgIC8vIEhSCiAgICBpZiAoL14tezMsfSQvLnRlc3QocykpIHsKICAgICAgY2hpbGRyZW4ucHVzaChuZXcgUGFyYWdyYXBoKHsKICAgICAgICBib3JkZXI6ICB7IGJvdHRvbTogeyBzdHlsZTogQm9yZGVyU3R5bGUuU0lOR0xFLCBzaXplOiA2LCBjb2xvcjogQy5taWRCbHVlIH0gfSwKICAgICAgICBzcGFjaW5nOiB7IGJlZm9yZTogMTIwLCBhZnRlcjogMTIwIH0sCiAgICAgICAgY2hpbGRyZW46IFtdLAogICAgICB9KSk7CiAgICAgIGkrKzsgY29udGludWU7CiAgICB9CgogICAgLy8g67mIIOykhAogICAgaWYgKCFzKSB7CiAgICAgIGNoaWxkcmVuLnB1c2gobmV3IFBhcmFncmFwaCh7IHNwYWNpbmc6IHsgYmVmb3JlOiAyMCwgYWZ0ZXI6IDIwIH0sIGNoaWxkcmVuOiBbXSB9KSk7CiAgICAgIGkrKzsgY29udGludWU7CiAgICB9CgogICAgLy8g7J2066+47KeAIOyehOuyoOuTnCDrp4jsu6QKICAgIGNvbnN0IGVtbSA9IHMubWF0Y2goL15cW0VNQkVEX0lNQUdFOiguKylcXSQvKTsKICAgIGlmIChlbW0pIHsKICAgICAgY29uc3QgaW1nUGFyYSA9IG1ha2VJbWFnZVBhcmFncmFwaChlbW1bMV0sIF9zb3VyY2VQYXRoIHx8ICIiKTsKICAgICAgaWYgKGltZ1BhcmEpIGNoaWxkcmVuLnB1c2goaW1nUGFyYSk7CiAgICAgIGkrKzsgY29udGludWU7CiAgICB9CgogICAgLy8g7J2867CYIOusuOuLqAogICAgY2hpbGRyZW4ucHVzaChuZXcgUGFyYWdyYXBoKHsgc3BhY2luZzogeyBiZWZvcmU6IDYwLCBhZnRlcjogNjAgfSwgY2hpbGRyZW46IHBhcnNlSW5saW5lKHMpIH0pKTsKICAgIGkrKzsKICB9CiAgZmx1c2hUYWJsZSgpOwogIHJldHVybiBjaGlsZHJlbjsKfQoKLy8g4pSA4pSA4pSAIOy7pOuyhCDtjpjsnbTsp4Ag4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSACmZ1bmN0aW9uIG1ha2VDb3Zlcih0aXRsZSwgc3VidGl0bGUsIGRhdGUpIHsKICByZXR1cm4gWwogICAgbmV3IFBhcmFncmFwaCh7IHNwYWNpbmc6IHsgYmVmb3JlOiAxODAwIH0sIGNoaWxkcmVuOiBbXSB9KSwKICAgIG5ldyBQYXJhZ3JhcGgoewogICAgICBib3JkZXI6IHsgdG9wOiB7IHN0eWxlOiBCb3JkZXJTdHlsZS5TSU5HTEUsIHNpemU6IDQwLCBjb2xvcjogQy5ibHVlIH0gfSwKICAgICAgc3BhY2luZzogeyBhZnRlcjogODAwIH0sIGNoaWxkcmVuOiBbXSwKICAgIH0pLAogICAgbmV3IFBhcmFncmFwaCh7CiAgICAgIGFsaWdubWVudDogQWxpZ25tZW50VHlwZS5DRU5URVIsIHNwYWNpbmc6IHsgYWZ0ZXI6IDI0MCB9LAogICAgICBjaGlsZHJlbjogW25ldyBUZXh0UnVuKHsgdGV4dDogdGl0bGUsIGJvbGQ6IHRydWUsIHNpemU6IDcyLCBjb2xvcjogQy5ibHVlLCBmb250OiAi66eR7J2AIOqzoOuUlSIgfSldLAogICAgfSksCiAgICBuZXcgUGFyYWdyYXBoKHsKICAgICAgYWxpZ25tZW50OiBBbGlnbm1lbnRUeXBlLkNFTlRFUiwgc3BhY2luZzogeyBhZnRlcjogNjAwIH0sCiAgICAgIGNoaWxkcmVuOiBbbmV3IFRleHRSdW4oeyB0ZXh0OiBzdWJ0aXRsZSwgc2l6ZTogMzIsIGNvbG9yOiBDLmdyYXksIGZvbnQ6ICLrp5HsnYAg6rOg65SVIiB9KV0sCiAgICB9KSwKICAgIG5ldyBQYXJhZ3JhcGgoewogICAgICBib3JkZXI6IHsgYm90dG9tOiB7IHN0eWxlOiBCb3JkZXJTdHlsZS5TSU5HTEUsIHNpemU6IDgsIGNvbG9yOiBDLm1pZEJsdWUgfSB9LAogICAgICBzcGFjaW5nOiB7IGFmdGVyOiA0MDAgfSwgY2hpbGRyZW46IFtdLAogICAgfSksCiAgICBuZXcgUGFyYWdyYXBoKHsKICAgICAgYWxpZ25tZW50OiBBbGlnbm1lbnRUeXBlLkNFTlRFUiwgc3BhY2luZzogeyBhZnRlcjogMjAwIH0sCiAgICAgIGNoaWxkcmVuOiBbbmV3IFRleHRSdW4oeyB0ZXh0OiBkYXRlLCBzaXplOiAyNiwgY29sb3I6IEMuZ3JheSwgZm9udDogIuunkeydgCDqs6DrlJUiIH0pXSwKICAgIH0pLAogICAgbmV3IFBhcmFncmFwaCh7CiAgICAgIGFsaWdubWVudDogQWxpZ25tZW50VHlwZS5DRU5URVIsCiAgICAgIGNoaWxkcmVuOiBbbmV3IFRleHRSdW4oeyB0ZXh0OiAiU2Ftc3VuZyBDb25maWRlbnRpYWwiLCBzaXplOiAyMiwgY29sb3I6IEMucmVkLCBib2xkOiB0cnVlLCBmb250OiAi66eR7J2AIOqzoOuUlSIgfSldLAogICAgfSksCiAgICBuZXcgUGFyYWdyYXBoKHsgY2hpbGRyZW46IFtuZXcgUGFnZUJyZWFrKCldIH0pLAogIF07Cn0KCi8vIOKUgOKUgOKUgCDtl6TrjZQgLyDtkbjthLAg4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSACmZ1bmN0aW9uIG1ha2VIZWFkZXIodGl0bGUpIHsKICByZXR1cm4gbmV3IEhlYWRlcih7IGNoaWxkcmVuOiBbbmV3IFBhcmFncmFwaCh7CiAgICBib3JkZXI6ICB7IGJvdHRvbTogeyBzdHlsZTogQm9yZGVyU3R5bGUuU0lOR0xFLCBzaXplOiA0LCBjb2xvcjogQy5ibHVlLCBzcGFjZTogMSB9IH0sCiAgICBzcGFjaW5nOiB7IGFmdGVyOiAxMDAgfSwKICAgIHRhYlN0b3BzOiBbeyB0eXBlOiAicmlnaHQiLCBwb3NpdGlvbjogOTAyNiB9XSwKICAgIGNoaWxkcmVuOiBbCiAgICAgIG5ldyBUZXh0UnVuKHsgdGV4dDogdGl0bGUsICAgICAgICAgICAgICAgICAgIHNpemU6IDE4LCBmb250OiAi66eR7J2AIOqzoOuUlSIsIGNvbG9yOiBDLmJsdWUsIGJvbGQ6IHRydWUgfSksCiAgICAgIG5ldyBUZXh0UnVuKHsgdGV4dDogIlx0U2Ftc3VuZyBDb25maWRlbnRpYWwiLCBzaXplOiAxOCwgZm9udDogIuunkeydgCDqs6DrlJUiLCBjb2xvcjogQy5ncmF5IH0pLAogICAgXSwKICB9KV0gfSk7Cn0KCmZ1bmN0aW9uIG1ha2VGb290ZXIoZGF0ZSkgewogIHJldHVybiBuZXcgRm9vdGVyKHsgY2hpbGRyZW46IFtuZXcgUGFyYWdyYXBoKHsKICAgIGJvcmRlcjogICAgeyB0b3A6IHsgc3R5bGU6IEJvcmRlclN0eWxlLlNJTkdMRSwgc2l6ZTogNCwgY29sb3I6IEMubWlkQmx1ZSwgc3BhY2U6IDEgfSB9LAogICAgc3BhY2luZzogICB7IGJlZm9yZTogODAgfSwKICAgIGFsaWdubWVudDogQWxpZ25tZW50VHlwZS5DRU5URVIsCiAgICBjaGlsZHJlbjogWwogICAgICBuZXcgVGV4dFJ1bih7IHRleHQ6IGAke2RhdGV9ICDCtyAgYCwgc2l6ZTogMTgsIGZvbnQ6ICLrp5HsnYAg6rOg65SVIiwgY29sb3I6IEMuZ3JheSB9KSwKICAgICAgbmV3IFRleHRSdW4oeyBjaGlsZHJlbjogW1BhZ2VOdW1iZXIuQ1VSUkVOVF0sICAgICBzaXplOiAxOCwgZm9udDogIuunkeydgCDqs6DrlJUiLCBjb2xvcjogQy5ncmF5IH0pLAogICAgICBuZXcgVGV4dFJ1bih7IHRleHQ6ICIgLyAiLCAgICAgICAgICAgICAgICAgICAgICAgICBzaXplOiAxOCwgZm9udDogIuunkeydgCDqs6DrlJUiLCBjb2xvcjogQy5ncmF5IH0pLAogICAgICBuZXcgVGV4dFJ1bih7IGNoaWxkcmVuOiBbUGFnZU51bWJlci5UT1RBTF9QQUdFU10sIHNpemU6IDE4LCBmb250OiAi66eR7J2AIOqzoOuUlSIsIGNvbG9yOiBDLmdyYXkgfSksCiAgICBdLAogIH0pXSB9KTsKfQoKLy8g4pSA4pSA4pSAIO2OmOydtOyngCDshKTsoJUg4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSACmNvbnN0IFBBR0VfUFJPUFMgPSB7CiAgcGFnZTogewogICAgc2l6ZTogICB7IHdpZHRoOiAxMTkwNiwgaGVpZ2h0OiAxNjgzOCB9LCAgICAgICAgICAvLyBBNAogICAgbWFyZ2luOiB7IHRvcDogMTQ0MCwgcmlnaHQ6IDE0NDAsIGJvdHRvbTogMTQ0MCwgbGVmdDogMTgwMCB9LAogIH0sCn07CgovLyDilIDilIDilIAg66y47IScIOyKpO2DgOydvCDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIAKZnVuY3Rpb24gZG9jU3R5bGVzKCkgewogIHJldHVybiB7CiAgICBkZWZhdWx0OiB7IGRvY3VtZW50OiB7IHJ1bjogeyBmb250OiAi66eR7J2AIOqzoOuUlSIsIHNpemU6IDIyIH0gfSB9LAogICAgcGFyYWdyYXBoU3R5bGVzOiBbCiAgICAgIHsgaWQ6ICJIZWFkaW5nMSIsIG5hbWU6ICJIZWFkaW5nIDEiLCBiYXNlZE9uOiAiTm9ybWFsIiwgbmV4dDogIk5vcm1hbCIsIHF1aWNrRm9ybWF0OiB0cnVlLAogICAgICAgIHJ1bjogeyBzaXplOiAzNiwgYm9sZDogdHJ1ZSwgZm9udDogIuunkeydgCDqs6DrlJUiLCBjb2xvcjogQy5ibHVlIH0sCiAgICAgICAgcGFyYWdyYXBoOiB7IHNwYWNpbmc6IHsgYmVmb3JlOiA0MDAsIGFmdGVyOiAyMDAgfSwgb3V0bGluZUxldmVsOiAwIH0gfSwKICAgICAgeyBpZDogIkhlYWRpbmcyIiwgbmFtZTogIkhlYWRpbmcgMiIsIGJhc2VkT246ICJOb3JtYWwiLCBuZXh0OiAiTm9ybWFsIiwgcXVpY2tGb3JtYXQ6IHRydWUsCiAgICAgICAgcnVuOiB7IHNpemU6IDMwLCBib2xkOiB0cnVlLCBmb250OiAi66eR7J2AIOqzoOuUlSIsIGNvbG9yOiBDLmJsdWUgfSwKICAgICAgICBwYXJhZ3JhcGg6IHsgc3BhY2luZzogeyBiZWZvcmU6IDMwMCwgYWZ0ZXI6IDE2MCB9LCBvdXRsaW5lTGV2ZWw6IDEgfSB9LAogICAgICB7IGlkOiAiSGVhZGluZzMiLCBuYW1lOiAiSGVhZGluZyAzIiwgYmFzZWRPbjogIk5vcm1hbCIsIG5leHQ6ICJOb3JtYWwiLCBxdWlja0Zvcm1hdDogdHJ1ZSwKICAgICAgICBydW46IHsgc2l6ZTogMjYsIGJvbGQ6IHRydWUsIGZvbnQ6ICLrp5HsnYAg6rOg65SVIiwgY29sb3I6ICIzMzMzMzMiIH0sCiAgICAgICAgcGFyYWdyYXBoOiB7IHNwYWNpbmc6IHsgYmVmb3JlOiAyNDAsIGFmdGVyOiAxMjAgfSwgb3V0bGluZUxldmVsOiAyIH0gfSwKICAgIF0sCiAgfTsKfQoKLy8g4pSA4pSA4pSAIOydtOuvuOyngCDsgr3snoUg4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSA4pSACmZ1bmN0aW9uIG1ha2VJbWFnZVBhcmFncmFwaChpbWdQYXRoLCBiYXNlUGF0aCkgewogIHRyeSB7CiAgICAvLyBpbWdQYXRoOiAiaW1hZ2VzL2ltZ18wMDEucG5nIiAo7IOB64yA6rK966GcKQogICAgY29uc3QgZnVsbFBhdGggPSBwYXRoLmlzQWJzb2x1dGUoaW1nUGF0aCkKICAgICAgPyBpbWdQYXRoCiAgICAgIDogcGF0aC5qb2luKHBhdGguZGlybmFtZShiYXNlUGF0aCksIGltZ1BhdGgpOwogICAgaWYgKCFmcy5leGlzdHNTeW5jKGZ1bGxQYXRoKSkgcmV0dXJuIG51bGw7CgogICAgY29uc3QgaW1nRGF0YSAgPSBmcy5yZWFkRmlsZVN5bmMoZnVsbFBhdGgpOwogICAgY29uc3QgZXh0ICAgICAgPSBwYXRoLmV4dG5hbWUoZnVsbFBhdGgpLnRvTG93ZXJDYXNlKCkucmVwbGFjZSgiLiIsICIiKTsKICAgIGNvbnN0IHR5cGVNYXAgID0geyBqcGc6ICJqcGciLCBqcGVnOiAianBnIiwgcG5nOiAicG5nIiwgZ2lmOiAiZ2lmIiwgYm1wOiAiYm1wIiwgd2VicDogInBuZyIgfTsKICAgIGNvbnN0IGltZ1R5cGUgID0gdHlwZU1hcFtleHRdIHx8ICJwbmciOwoKICAgIC8vIOybkOuzuCDtgazquLAg7LaU7KCVIOyXhuydtCDstZzrjIAg7Y+tIOq4sOykgOycvOuhnCDsgr3snoUgKEE0IGNvbnRlbnQg7Y+tIOq4sOykgCkKICAgIGNvbnN0IG1heFcgPSA1MDAsIG1heEggPSA0MDA7CiAgICByZXR1cm4gbmV3IFBhcmFncmFwaCh7CiAgICAgIGFsaWdubWVudDogQWxpZ25tZW50VHlwZS5DRU5URVIsCiAgICAgIHNwYWNpbmc6IHsgYmVmb3JlOiAxMjAsIGFmdGVyOiAxMjAgfSwKICAgICAgY2hpbGRyZW46IFtuZXcgSW1hZ2VSdW4oewogICAgICAgIHR5cGU6IGltZ1R5cGUsCiAgICAgICAgZGF0YTogaW1nRGF0YSwKICAgICAgICB0cmFuc2Zvcm1hdGlvbjogeyB3aWR0aDogbWF4VywgaGVpZ2h0OiBtYXhIIH0sCiAgICAgICAgYWx0VGV4dDogeyB0aXRsZTogIuydtOuvuOyngCIsIGRlc2NyaXB0aW9uOiBpbWdQYXRoLCBuYW1lOiBpbWdQYXRoIH0sCiAgICAgIH0pXSwKICAgIH0pOwogIH0gY2F0Y2ggKGUpIHsKICAgIHJldHVybiBudWxsOwogIH0KfQoKLy8g4pSA4pSA4pSAIOyGjOyKpCDroZzrk5wgKO2MjOydvCBvciDtj7TrjZQpIOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgOKUgApmdW5jdGlvbiBsb2FkU291cmNlcyhzcmMpIHsKICBjb25zdCBzdGF0ID0gZnMuc3RhdFN5bmMoc3JjKTsKICBpZiAoc3RhdC5pc0ZpbGUoKSkgewogICAgcmV0dXJuIFt7IHRpdGxlOiBwYXRoLmJhc2VuYW1lKHNyYywgIi5tZCIpLCBtZDogZnMucmVhZEZpbGVTeW5jKHNyYywgInV0ZjgiKSwgcGF0aDogc3JjIH1dOwogIH0KICAvLyDtj7TrjZQ6IC5tZCDtjIzsnbwg7KCV66CsIOuhnOuTnCAo67O06rOg7IScL3JlcG9ydCDtjIzsnbwg7KCc7Jm4KQogIHJldHVybiBmcy5yZWFkZGlyU3luYyhzcmMpCiAgICAuZmlsdGVyKGYgPT4gZi5lbmRzV2l0aCgiLm1kIikgJiYgIS/rs7Tqs6DshJx8cmVwb3J0L2kudGVzdChmKSkKICAgIC5zb3J0KCkKICAgIC5tYXAoZiA9PiAoewogICAgICB0aXRsZTogZi5yZXBsYWNlKCIubWQiLCAiIiksCiAgICAgIG1kOiAgICBmcy5yZWFkRmlsZVN5bmMocGF0aC5qb2luKHNyYywgZiksICJ1dGY4IiksCiAgICAgIHBhdGg6ICBwYXRoLmpvaW4oc3JjLCBmKSwKICAgIH0pKTsKfQoKLy8g4pSA4pSA4pSAIOuplOyduCDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIDilIAKYXN5bmMgZnVuY3Rpb24gbWFpbigpIHsKICBjb25zdCBbLCwgc3JjLCBvdXRQYXRoLCB0aXRsZUFyZ10gPSBwcm9jZXNzLmFyZ3Y7CiAgaWYgKCFzcmMgfHwgIW91dFBhdGgpIHsKICAgIGNvbnNvbGUuZXJyb3IoIlVzYWdlOiBub2RlIG1ha2VfcmVwb3J0LmpzIDxmaWxlLm1kfGZvbGRlcj4gPG91dHB1dC5kb2N4PiBbdGl0bGVdIik7CiAgICBwcm9jZXNzLmV4aXQoMSk7CiAgfQoKICBjb25zdCBzb3VyY2VzICA9IGxvYWRTb3VyY2VzKHNyYyk7CiAgY29uc3QgdGl0bGUgICAgPSB0aXRsZUFyZyB8fCBwYXRoLmJhc2VuYW1lKHNyYywgIi5tZCIpIHx8ICJDb25mbHVlbmNlIOuztOqzoOyEnCI7CiAgY29uc3QgZGF0ZSAgICAgPSBuZXcgRGF0ZSgpLnRvTG9jYWxlRGF0ZVN0cmluZygia28tS1IiLCB7IHllYXI6ICJudW1lcmljIiwgbW9udGg6ICJsb25nIiwgZGF5OiAibnVtZXJpYyIgfSk7CiAgY29uc3Qgc3VidGl0bGUgPSBgQ29uZmx1ZW5jZSDsnpDrj5kg7IiY7KeRIOuztOqzoOyEnCAgwrcgIOy0nSAke3NvdXJjZXMubGVuZ3RofeqwnCDtjpjsnbTsp4BgOwoKICBjb25zb2xlLmxvZyhg8J+ThCDshozsiqQgJHtzb3VyY2VzLmxlbmd0aH3qsJwg66Gc65OcIOyZhOujjGApOwoKICAvLyDrs7jrrLggY2hpbGRyZW4g6rWs7ISxCiAgY29uc3QgYm9keUNoaWxkcmVuID0gW107CiAgc291cmNlcy5mb3JFYWNoKChzLCBpZHgpID0+IHsKICAgIGlmIChpZHggPiAwKSBib2R5Q2hpbGRyZW4ucHVzaChuZXcgUGFyYWdyYXBoKHsgY2hpbGRyZW46IFtuZXcgUGFnZUJyZWFrKCldIH0pKTsKICAgIGlmIChzb3VyY2VzLmxlbmd0aCA+IDEpIGJvZHlDaGlsZHJlbi5wdXNoKG1ha2VTb3VyY2VCYW5uZXIocy50aXRsZSkpOwogICAgYm9keUNoaWxkcmVuLnB1c2goLi4ucGFyc2VNRChzLm1kLCBzLnBhdGgpKTsKICB9KTsKCiAgY29uc3QgZG9jID0gbmV3IERvY3VtZW50KHsKICAgIHN0eWxlczogICBkb2NTdHlsZXMoKSwKICAgIHNlY3Rpb25zOiBbCiAgICAgIC8vIOKUgOKUgCDshLnshZgxOiDsu6TrsoQgKyDrqqnssKgg4pSA4pSACiAgICAgIHsKICAgICAgICBwcm9wZXJ0aWVzOiB7IHBhZ2U6IFBBR0VfUFJPUFMucGFnZSB9LAogICAgICAgIGNoaWxkcmVuOiBbCiAgICAgICAgICAuLi5tYWtlQ292ZXIodGl0bGUsIHN1YnRpdGxlLCBkYXRlKSwKICAgICAgICAgIG1ha2VIZWFkaW5nKCLrqqkgIOywqCIsIDEpLAogICAgICAgICAgbmV3IFRhYmxlT2ZDb250ZW50cygi66qp7LCoIiwgeyBoeXBlcmxpbms6IHRydWUsIGhlYWRpbmdTdHlsZVJhbmdlOiAiMS0zIiB9KSwKICAgICAgICAgIG5ldyBQYXJhZ3JhcGgoeyBjaGlsZHJlbjogW25ldyBQYWdlQnJlYWsoKV0gfSksCiAgICAgICAgXSwKICAgICAgfSwKICAgICAgLy8g7IS57IWYMjog67O466y4CiAgICAgIHsKICAgICAgICBwcm9wZXJ0aWVzOiB7IHBhZ2U6IFBBR0VfUFJPUFMucGFnZSB9LAogICAgICAgIGhlYWRlcnM6ICAgIHsgZGVmYXVsdDogbWFrZUhlYWRlcih0aXRsZSkgIH0sCiAgICAgICAgZm9vdGVyczogICAgeyBkZWZhdWx0OiBtYWtlRm9vdGVyKGRhdGUpICAgfSwKICAgICAgICBjaGlsZHJlbjogICBib2R5Q2hpbGRyZW4sCiAgICAgIH0sCiAgICBdLAogIH0pOwoKICBjb25zdCBidWYgPSBhd2FpdCBQYWNrZXIudG9CdWZmZXIoZG9jKTsKICBmcy53cml0ZUZpbGVTeW5jKG91dFBhdGgsIGJ1Zik7CiAgY29uc29sZS5sb2coYOKchSBXb3JkIOuztOqzoOyEnCDsg53shLEg7JmE66OMOiAke291dFBhdGh9YCk7Cn0KCm1haW4oKS5jYXRjaChlID0+IHsgY29uc29sZS5lcnJvcigi4p2MIiwgZS5tZXNzYWdlKTsgcHJvY2Vzcy5leGl0KDEpOyB9KTsK"

# ─────────────────────────────────────────────────────────────────────────────
# Samsung-style Word report generator (python-docx, no Node.js required)
# ─────────────────────────────────────────────────────────────────────────────
def _make_samsung_word(md_sources: list, out_docx: str, title: str = "", callback=None) -> bool:
    """
    md_sources: [{"title": str, "md": str, "path": str}, ...]
    Generates a Samsung Blue styled .docx directly with python-docx.
    Returns True on success.
    """
    try:
        docx_mod   = _import_or_install("docx", "python-docx")
        Document   = docx_mod.Document
        from docx.shared    import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns   import qn
        from docx.oxml      import OxmlElement


        BLUE   = RGBColor(0x14, 0x28, 0xA0)
        BLUE_H = "1428A0"
        LGRAY  = RGBColor(0xF4, 0xF6, 0xFD)
        WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
        DGRAY  = RGBColor(0x44, 0x44, 0x44)
        RED    = RGBColor(0xDC, 0x26, 0x26)

        doc = Document()
        # Page setup A4
        for sec in doc.sections:
            sec.page_width  = Cm(21);  sec.page_height = Cm(29.7)
            sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
            sec.top_margin  = Cm(2.5); sec.bottom_margin = Cm(2.5)

        # ── Global font ──────────────────────────────────────────────────────
        style = doc.styles["Normal"]
        style.font.name = "맑은 고딕"; style.font.size = Pt(11)
        for hn, sz in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
            if hn in doc.styles:
                h = doc.styles[hn]
                h.font.name = "맑은 고딕"; h.font.size = Pt(sz)
                h.font.color.rgb = BLUE; h.font.bold = True

        # ── Helper: set cell bg color ────────────────────────────────────────
        def cell_bg(cell, hex_color):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:shd")):
                tcPr.remove(old)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  hex_color)
            tcPr.append(shd)

        # ── Helper: add a blue HR paragraph ─────────────────────────────────
        def add_hr():
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), BLUE_H)
            pBdr.append(bot)
            pPr.append(pBdr)
            return p

        # ── Cover page ───────────────────────────────────────────────────────
        if callback: callback("[Word] 커버 페이지 생성 중...")
        cov = doc.add_paragraph(); cov.paragraph_format.space_before = Pt(72)
        add_hr()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title or "Confluence 보고서")
        r.font.name = "맑은 고딕"; r.font.size = Pt(28); r.font.bold = True
        r.font.color.rgb = BLUE
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
        r2.font.name = "맑은 고딕"; r2.font.size = Pt(12); r2.font.color.rgb = DGRAY
        doc.add_paragraph()
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run("Samsung Confidential")
        r3.font.name = "맑은 고딕"; r3.font.size = Pt(11); r3.font.bold = True
        r3.font.color.rgb = RED
        add_hr()
        doc.add_page_break()

        # ── Body: parse each MD source ───────────────────────────────────────
        for si, src in enumerate(md_sources):
            if si > 0:
                doc.add_page_break()
            if len(md_sources) > 1:
                banner = doc.add_paragraph()
                banner.paragraph_format.space_before = Pt(4)
                banner.paragraph_format.space_after  = Pt(8)
                br = banner.add_run(f"📄  {src['title']}")
                br.font.name = "맑은 고딕"; br.font.size = Pt(10); br.font.color.rgb = RGBColor(0x6B,0x72,0x80)

            if callback: callback(f"[Word] 변환 중: {src['title']}")
            base_dir = os.path.dirname(os.path.abspath(src["path"])) if src.get("path") else ""
            _parse_md_to_doc(doc, src["md"], base_dir,
                             BLUE, BLUE_H, WHITE, LGRAY, DGRAY, RED,
                             cell_bg, Pt, Cm, RGBColor, WD_ALIGN_PARAGRAPH,
                             callback=callback)

        # ── Header / Footer ──────────────────────────────────────────────────
        report_title = title or "Confluence 보고서"
        today_str = datetime.now().strftime("%Y-%m-%d")
        for sec in doc.sections:
            # ── Header (simple text, no XML manipulation) ──
            hdr = sec.header
            hp  = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
            hp.clear()
            run_h = hp.add_run(f"{report_title}    |    Samsung Confidential")
            run_h.font.name = "맑은 고딕"
            run_h.font.size = Pt(9)
            run_h.font.color.rgb = DGRAY
            # ── Footer: date + page field (each field token in its own w:r) ──
            ftr = sec.footer
            fp  = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
            fp.clear()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_d = fp.add_run(f"{today_str}   ")
            run_d.font.name = "맑은 고딕"
            run_d.font.size = Pt(9)
            run_d.font.color.rgb = DGRAY
            # Page number via field codes — each token must live in its own <w:r>
            def _add_field_run(para, fchar_type=None, instr_text=None):
                r_el = OxmlElement("w:r")
                if fchar_type:
                    fc = OxmlElement("w:fldChar")
                    fc.set(qn("w:fldCharType"), fchar_type)
                    r_el.append(fc)
                else:
                    it = OxmlElement("w:instrText")
                    it.set(qn("xml:space"), "preserve")
                    it.text = instr_text
                    r_el.append(it)
                para._p.append(r_el)
            _add_field_run(fp, fchar_type="begin")
            _add_field_run(fp, instr_text=" PAGE ")
            _add_field_run(fp, fchar_type="end")

        doc.save(out_docx)
        if callback: callback(f"[Word] ✅ 완료 → {os.path.basename(out_docx)}")
        return True
    except Exception as e:
        import traceback
        if callback: callback(f"[Word] ❌ 오류: {e}\n{traceback.format_exc()[:300]}")
        return False


def _parse_md_to_doc(doc, md_text, base_dir,
                     BLUE, BLUE_H, WHITE, LGRAY, DGRAY, RED,
                     cell_bg, Pt, Cm, RGBColor, WD_ALIGN_PARAGRAPH,
                     callback=None):
    from docx.oxml.ns import qn
    from docx.oxml    import OxmlElement
    # Ensure Pillow is available for add_picture
    try:
        import PIL
    except ImportError:
        import subprocess as _sp, sys as _sys
        _sp.run([_sys.executable, "-m", "pip", "install", "Pillow", "--quiet"], check=False)

    in_code = False; code_lines = []
    table_rows = []; in_table = False

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows: in_table = False; return
        try:
            headers   = table_rows[0]
            data_rows = [r for r in table_rows[1:]
                         if not all(re.match(r"^[-:]+$", c.strip()) for c in r)]
            if not data_rows: table_rows.clear(); in_table = False; return
            ncols = max(len(headers), max((len(r) for r in data_rows), default=1))
            tbl = doc.add_table(rows=1 + len(data_rows), cols=ncols)
            tbl.style = "Table Grid"
            # (table alignment omitted for compatibility)
            # Header row
            for ci, h in enumerate(headers[:ncols]):
                c = tbl.rows[0].cells[ci]; c.text = h.strip()
                cell_bg(c, BLUE_H)
                for run in c.paragraphs[0].runs:
                    run.bold = True; run.font.color.rgb = WHITE
                    run.font.name = "맑은 고딕"; run.font.size = Pt(10)
            # Data rows
            for ri, row in enumerate(data_rows):
                fill = "F4F6FD" if ri % 2 == 0 else "FFFFFF"
                for ci, val in enumerate(row[:ncols]):
                    c = tbl.rows[ri+1].cells[ci]; c.text = val.strip()
                    cell_bg(c, fill)
                    for run in c.paragraphs[0].runs:
                        run.font.name = "맑은 고딕"; run.font.size = Pt(10)
            doc.add_paragraph()
        except Exception as _te:
            if callback: callback(f"[Word] 표 생성 오류: {_te}")
        table_rows.clear(); in_table = False

    def embed_image(path_str):
        # Normalize path: support both / and \ separators
        norm = path_str.replace("/", os.sep).replace("\\", os.sep)
        if os.path.isabs(norm):
            img_abs = norm
        elif base_dir:
            img_abs = os.path.normpath(os.path.join(base_dir, norm))
        else:
            img_abs = os.path.normpath(norm)
        if callback: callback(f"[Word] 이미지 삽입 시도: {img_abs}")
        if not os.path.exists(img_abs):
            if callback: callback(f"[Word] ⚠️ 이미지 없음: {img_abs}")
            return
        try:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_abs, width=Cm(14))
            doc.add_paragraph()
            if callback: callback(f"[Word] ✅ 이미지 삽입 완료: {os.path.basename(img_abs)}")
        except Exception as _e:
            if callback: callback(f"[Word] ❌ 이미지 삽입 실패: {_e}")
            import traceback as _tb
            if callback: callback(_tb.format_exc()[:400])

    for line in md_text.splitlines():
        s = line.strip()
        # Code block
        if s.startswith("```"):
            if in_table: flush_table()
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                r = p.add_run("\n".join(code_lines))
                r.font.name = "Consolas"; r.font.size = Pt(9)
                code_lines.clear(); in_code = False
            else:
                in_code = True
            continue
        if in_code: code_lines.append(line); continue
        # Table
        if s.startswith("|") and "|" in s[1:]:
            cols = [c.strip() for c in s.split("|")[1:-1]]
            table_rows.append(cols); in_table = True; continue
        elif in_table: flush_table()
        # EMBED_IMAGE
        if s.startswith("[EMBED_IMAGE:") and s.endswith("]"):
            raw_path = s[len("[EMBED_IMAGE:"):-1]
            if callback: callback(f"[Word] EMBED_IMAGE 마커: {raw_path}")
            embed_image(raw_path); continue
        # Headings
        hm = re.match(r"^(#{1,4})\s+(.*)", s)
        if hm:
            level = min(len(hm.group(1)), 4)
            p = doc.add_heading(hm.group(2), level=level)
            for run in p.runs:
                run.font.name = "맑은 고딕"; run.font.color.rgb = BLUE; run.bold = True
            continue
        # Bullet
        bm = re.match(r"^[-*•]\s+(.*)", s)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(bm.group(1))
            r.font.name = "맑은 고딕"; r.font.size = Pt(11)
            continue
        # Numbered
        nm = re.match(r"^\d+\.\s+(.*)", s)
        if nm:
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(nm.group(1))
            r.font.name = "맑은 고딕"; r.font.size = Pt(11)
            continue
        # Blockquote / callout
        qm = re.match(r"^>\s*(.*)", s)
        if qm:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(qm.group(1))
            r.font.name = "맑은 고딕"; r.font.size = Pt(10); r.italic = True
            r.font.color.rgb = RGBColor(0x2D, 0x3A, 0x8A)
            continue
        # Vision callout lines
        if s.startswith("**[이미지") or s.startswith("**[Vision"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            clean = re.sub(r"\*\*\[.*?\]\*\*\s*", "", s).strip("*").strip()
            r = p.add_run("📷  " + clean)
            r.font.name = "맑은 고딕"; r.font.size = Pt(10); r.italic = True
            r.font.color.rgb = RGBColor(0x2D, 0x3A, 0x8A)
            continue
        # HR
        if re.match(r"^-{3,}$", s):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single"); bot.set(qn("w:sz"),    "4")
            bot.set(qn("w:space"), "1");      bot.set(qn("w:color"), "C7CDE8")
            pBdr.append(bot); pPr.append(pBdr)
            continue
        # Empty
        if not s: doc.add_paragraph(); continue
        # Regular text (strip inline ** * `)
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
        clean = re.sub(r"\*(.+?)\*",     r"\1", clean)
        clean = re.sub(r"`(.+?)`",       r"\1", clean)
        p = doc.add_paragraph()
        r = p.add_run(clean)
        r.font.name = "맑은 고딕"; r.font.size = Pt(11)
    if in_table: flush_table()


def generate_word_report(src: str, out_docx: str, title: str = "", callback=None) -> str:
    """Generate Word report from an MD file or folder. Returns out_docx path on success."""
    src_abs     = os.path.abspath(src)
    out_docx_abs = os.path.abspath(out_docx)

    # Build source list
    if os.path.isfile(src_abs):
        sources = [{"title": os.path.basename(src_abs).replace(".md",""),
                    "md":    open(src_abs, encoding="utf-8").read(),
                    "path":  src_abs}]
    else:
        mds = sorted(f for f in os.listdir(src_abs) if f.endswith(".md")
                     and not re.search(r"보고서|report", f, re.I))
        sources = [{"title": f.replace(".md",""),
                    "md":    open(os.path.join(src_abs,f), encoding="utf-8").read(),
                    "path":  os.path.join(src_abs,f)} for f in mds]

    if not sources:
        if callback: callback("[Word] 변환할 MD 파일 없음")
        return ""

    ok = _make_samsung_word(sources, out_docx_abs, title=title, callback=callback)

    # Optional: try Node.js for even better quality (TOC etc.) - silent fallback
    if not ok and _shutil.which("node"):
        try:
            ok = _generate_word_node(src_abs, out_docx_abs, title, callback)
        except Exception:
            pass

    return out_docx_abs if ok and os.path.exists(out_docx_abs) else ""


def _generate_word_node(src, out_docx, title="", callback=None):
    """Try Node.js path (optional, for TOC support)."""
    work_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_word_engine")
    os.makedirs(work_dir, exist_ok=True)
    js_path = os.path.join(work_dir, "make_report.js")
    js_txt  = _base64.b64decode(_MAKE_REPORT_JS_B64).decode("utf-8")
    try:
        existing = open(js_path, encoding="utf-8").read() if os.path.exists(js_path) else ""
    except Exception:
        existing = ""
    if existing != js_txt:
        with open(js_path, "w", encoding="utf-8") as f: f.write(js_txt)
    nm = os.path.join(work_dir, "node_modules", "docx")
    if not os.path.isdir(nm):
        r = _subprocess.run(["npm","install","docx"], cwd=work_dir, capture_output=True, timeout=120)
        if not os.path.isdir(nm): return False
    cmd = ["node", js_path, src, out_docx] + ([title] if title else [])
    r = _subprocess.run(cmd, capture_output=True, text=True, timeout=120, cwd=work_dir)
    if r.returncode != 0:
        if callback: callback(f"[Word-Node] {(r.stderr or r.stdout or '').strip()[:200]}")
        return False
    if callback: callback("[Word] ✅ Node.js 보고서 완료 (TOC 포함)")
    return True


def setup_style():
    style = ttk.Style()
    style.theme_use("clam")
    style.configure(".", background=BG, foreground=TEXT, font=("Malgun Gothic", 10))
    style.configure("Card.TFrame", background=CARD, relief="flat")
    style.configure("TLabel", background=CARD, foreground=TEXT, font=("Malgun Gothic", 10))
    style.configure("Root.TFrame", background=BG)
    style.configure("Title.TLabel", background=CARD, font=("Malgun Gothic", 13, "bold"), foreground=BLUE_DK)
    style.configure("Muted.TLabel", background=CARD, foreground=TEXT_MUTED)
    style.configure("TButton", background=BLUE, foreground="white", font=("Malgun Gothic", 10, "bold"), padding=(14, 7))
    style.map("TButton", background=[("active", BLUE_DK), ("disabled", BORDER)])
    style.configure("TEntry", fieldbackground="white", foreground=TEXT, padding=(8, 6))
    style.configure("TCheckbutton", background=CARD, foreground=TEXT)
    style.configure("TProgressbar", background=BLUE, troughcolor=BORDER)
    style.configure("TNotebook", background=BG, borderwidth=0)
    style.configure("TNotebook.Tab", background=BG, foreground=TEXT_MUTED, padding=[18, 8])
    style.map("TNotebook.Tab", background=[("selected", CARD)], foreground=[("selected", BLUE_DK)])


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Confluence MD 수집 + OCR 보고서 생성 v2")
        self.geometry("940x720")
        self.resizable(True, True)
        self.configure(bg=BG)
        setup_style()

        nb = ttk.Notebook(self)
        nb.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)
        self.tab_crawl = CrawlTab(nb)
        self.tab_report = ReportTab(nb)
        self.tab_settings = SettingsTab(nb)
        nb.add(self.tab_crawl, text="수집")
        nb.add(self.tab_report, text="보고서 생성")
        nb.add(self.tab_settings, text="LLM 설정")


class CrawlTab(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent, style="Card.TFrame", padding=20)
        self._build()

    def _build(self):
        self.columnconfigure(1, weight=1)
        self.rowconfigure(10, weight=1)

        ttk.Label(self, text="Confluence URL:", style="Title.TLabel").grid(row=0, column=0, sticky=tk.W, pady=8)
        self.url_var = tk.StringVar()
        ttk.Entry(self, textvariable=self.url_var, width=70).grid(row=0, column=1, columnspan=2, sticky=tk.EW, pady=8)
        self.url_var.trace_add("write", self._on_url)
        self.pid_lbl = ttk.Label(self, text="페이지 ID: -", style="Muted.TLabel")
        self.pid_lbl.grid(row=1, column=1, sticky=tk.W, pady=4)

        ttk.Label(self, text="재귀 깊이:").grid(row=2, column=0, sticky=tk.W, pady=8)
        self.depth_var = tk.IntVar(value=3)
        ttk.Spinbox(self, from_=0, to=10, textvariable=self.depth_var, width=8).grid(row=2, column=1, sticky=tk.W)

        self.vision_var = tk.BooleanVar(value=True)
        self.summary_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(self, text="이미지 Vision 분석", variable=self.vision_var).grid(row=3, column=1, sticky=tk.W, pady=4)
        ttk.Checkbutton(self, text="페이지별 LLM 요약 생성", variable=self.summary_var).grid(row=4, column=1, sticky=tk.W, pady=4)
        self.word_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(self, text="수집 완료 후 Word 보고서 자동 생성 (.docx)", variable=self.word_var).grid(row=5, column=1, sticky=tk.W, pady=4)

        ttk.Label(self, text="출력 폴더:").grid(row=6, column=0, sticky=tk.W, pady=8)
        self.out_var = tk.StringVar(value="./confluence_output")
        frm = ttk.Frame(self, style="Card.TFrame")
        frm.grid(row=5, column=1, columnspan=2, sticky=tk.EW)
        frm.columnconfigure(0, weight=1)
        ttk.Entry(frm, textvariable=self.out_var).grid(row=0, column=0, sticky=tk.EW)
        ttk.Button(frm, text="찾아보기", command=self._browse).grid(row=0, column=1, padx=8)

        self.btn = ttk.Button(self, text="수집 시작", command=self._run)
        self.btn.grid(row=7, column=0, pady=12)
        self.prog = ttk.Progressbar(self, mode="indeterminate", length=700)
        self.prog.grid(row=8, column=0, columnspan=3, sticky=tk.EW)

        ttk.Label(self, text="로그:", style="Title.TLabel").grid(row=9, column=0, sticky=tk.W, pady=8)
        lf = ttk.Frame(self, style="Card.TFrame")
        lf.grid(row=10, column=0, columnspan=3, sticky=tk.NSEW)
        sb = ttk.Scrollbar(lf)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_box = tk.Text(lf, height=16, yscrollcommand=sb.set, wrap=tk.WORD, bg=LOG_BG, fg=LOG_FG, font=("Consolas", 9))
        self.log_box.pack(fill=tk.BOTH, expand=True)
        sb.config(command=self.log_box.yview)

    def _on_url(self, *_):
        pid = extract_page_id_from_url(self.url_var.get())
        self.pid_lbl.config(text=f"페이지 ID: {pid}" if pid else "페이지 ID: 찾을 수 없음")

    def _browse(self):
        d = filedialog.askdirectory()
        if d:
            self.out_var.set(d)

    def log(self, msg):
        try:
            self.log_box.insert(tk.END, msg + "\n")
            self.log_box.see(tk.END)
            self.update_idletasks()
        except Exception:
            pass

    def _run(self):
        pid = extract_page_id_from_url(self.url_var.get())
        if not pid:
            messagebox.showerror("오류", "URL에서 페이지 ID를 찾을 수 없습니다.")
            return
        self.btn.config(state="disabled")
        self.prog.start()
        threading.Thread(target=self._worker, args=(pid,), daemon=True).start()

    def _worker(self, page_id):
        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            self.after(0, lambda: messagebox.showerror("오류", "playwright 미설치: pip install playwright 후 playwright install chromium 실행 필요"))
            self.after(0, lambda: (self.prog.stop(), self.btn.config(state="normal")))
            return

        save_dir = os.path.join(self.out_var.get(), f"page_{page_id}")
        os.makedirs(save_dir, exist_ok=True)
        try:
            with sync_playwright() as p:
                browser = p.chromium.launch_persistent_context(
                    user_data_dir=USER_DATA_DIR,
                    headless=False,
                    viewport={"width": 1280, "height": 720},
                )
                page = browser.new_page()
                page.goto(f"{BASE_URL}/pages/viewpage.action?pageId={page_id}")
                page.wait_for_load_state("networkidle", timeout=30000)

                if not messagebox.askokcancel("로그인 확인", "Confluence에 로그인되어 있나요?\n[확인] 진행 / [취소] 중단"):
                    browser.close()
                    return

                page.reload(wait_until="networkidle")
                self.log("=" * 50)
                self.log(f"수집 시작 | 페이지 ID: {page_id}")
                self.log(f"이미지 OCR: {self.vision_var.get()} | LLM 요약: {self.summary_var.get()}")
                self.log("=" * 50)

                process_page(
                    page_session=page,
                    page_id=page_id,
                    save_dir=save_dir,
                    depth=self.depth_var.get(),
                    use_vision=self.vision_var.get(),
                    use_llm_summary=self.summary_var.get(),
                    callback=self.log,
                )
                browser.close()
                self.log("\n수집 완료!")
                self.log(f"저장 위치: {os.path.abspath(save_dir)}")

                # Word 보고서 자동 생성
                if self.word_var.get():
                    page_title_raw = self.url_var.get().split("/")[-1].replace("+", " ")
                    docx_path = os.path.join(os.path.abspath(self.out_var.get()), f"보고서_{page_id}.docx")
                    generate_word_report(os.path.abspath(save_dir), docx_path,
                                         title=page_title_raw or f"Confluence 보고서",
                                         callback=self.log)
                    if os.path.exists(docx_path):
                        messagebox.showinfo("완료", f"수집 및 Word 보고서 생성 완료!\n\n📁 MD: {os.path.abspath(save_dir)}\n📄 Word: {os.path.abspath(docx_path)}")
                    else:
                        messagebox.showinfo("완료", f"수집 완료!\n{os.path.abspath(save_dir)}")
                else:
                    messagebox.showinfo("완료", f"수집 완료!\n{os.path.abspath(save_dir)}")
        except Exception as e:
            self.log(f"[오류] {e}\n{traceback.format_exc()}")
            messagebox.showerror("오류", str(e))
        finally:
            self.after(0, lambda: (self.prog.stop(), self.btn.config(state="normal")))


class ReportTab(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent, style="Card.TFrame", padding=20)
        self.md_files: List[str] = []
        self._build()

    def _build(self):
        self.columnconfigure(0, weight=1)
        self.rowconfigure(1, weight=1)
        self.rowconfigure(7, weight=1)

        # ── 상단 버튼 바 ──────────────────────────────────
        top = ttk.Frame(self, style="Card.TFrame")
        top.grid(row=0, column=0, columnspan=2, sticky=tk.EW, pady=(0, 6))
        ttk.Label(top, text="파일 관리:", style="Title.TLabel").pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(top, text="📂 폴더 추가",    command=self._load_folder).pack(side=tk.LEFT, padx=3)
        ttk.Button(top, text="📄 파일 추가",    command=self._load_files).pack(side=tk.LEFT, padx=3)
        ttk.Separator(top, orient="vertical").pack(side=tk.LEFT, fill=tk.Y, padx=6, pady=3)
        ttk.Button(top, text="✅ 전체 선택",    command=self._check_all).pack(side=tk.LEFT, padx=3)
        ttk.Button(top, text="☐ 전체 해제",    command=self._uncheck_all).pack(side=tk.LEFT, padx=3)
        ttk.Separator(top, orient="vertical").pack(side=tk.LEFT, fill=tk.Y, padx=6, pady=3)
        ttk.Button(top, text="🗑 선택 제거",    command=self._remove_selected).pack(side=tk.LEFT, padx=3)
        ttk.Button(top, text="🧹 목록 초기화",  command=self._clear_all).pack(side=tk.LEFT, padx=3)
        self.file_count_lbl = ttk.Label(top, text="0개 파일", style="Muted.TLabel")
        self.file_count_lbl.pack(side=tk.RIGHT, padx=8)

        # ── 파일 목록 (Treeview + 체크박스) ───────────────
        lf = ttk.LabelFrame(self, text="파일 목록  (행 클릭 = 체크 토글 / 체크된 파일만 보고서에 포함)")
        lf.grid(row=1, column=0, columnspan=2, sticky=tk.NSEW, pady=4)
        lf.columnconfigure(0, weight=1)
        lf.rowconfigure(0, weight=1)

        cols = ("check", "filename", "folder")
        self.file_list = ttk.Treeview(lf, columns=cols, show="headings", height=12,
                                       selectmode="extended")
        self.file_list.heading("check",    text="포함")
        self.file_list.heading("filename", text="파일명")
        self.file_list.heading("folder",   text="경로")
        self.file_list.column("check",    width=50,  minwidth=50,  stretch=False, anchor="center")
        self.file_list.column("filename", width=260, minwidth=120, stretch=True)
        self.file_list.column("folder",   width=320, minwidth=120, stretch=True)
        sb_y = ttk.Scrollbar(lf, orient="vertical",   command=self.file_list.yview)
        sb_x = ttk.Scrollbar(lf, orient="horizontal", command=self.file_list.xview)
        self.file_list.configure(yscrollcommand=sb_y.set, xscrollcommand=sb_x.set)
        self.file_list.grid(row=0, column=0, sticky=tk.NSEW)
        sb_y.grid(row=0, column=1, sticky=tk.NS)
        sb_x.grid(row=1, column=0, sticky=tk.EW)
        self.file_list.bind("<ButtonRelease-1>", self._on_tree_click)
        self.file_list.tag_configure("checked",   background="#EFF6FF")
        self.file_list.tag_configure("unchecked", background="#FFFFFF", foreground="#9CA3AF")

        # Word 변환 버튼 (파일 목록 바로 아래)
        word_bar = ttk.Frame(self, style="Card.TFrame")
        word_bar.grid(row=2, column=0, columnspan=2, sticky=tk.EW, pady=(2, 6))
        ttk.Button(word_bar, text="📝 체크된 파일 → Word 변환", command=self._export_to_word).pack(side=tk.LEFT, padx=4)
        ttk.Label(word_bar, text="(체크된 파일만 하나의 Word 보고서로 합칩니다)", style="Muted.TLabel").pack(side=tk.LEFT)

        opt = ttk.LabelFrame(self, text="보고서 옵션")
        opt.grid(row=3, column=0, columnspan=2, sticky=tk.EW, pady=10)
        self.report_title_var = tk.StringVar(value="주간 보고서")
        _desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        _default_out = _desktop if os.path.isdir(_desktop) else os.path.expanduser("~")
        self.out_var = tk.StringVar(value=_default_out)
        ttk.Label(opt, text="보고서 제목:").grid(row=0, column=0, padx=6, pady=6, sticky=tk.W)
        ttk.Entry(opt, textvariable=self.report_title_var, width=40).grid(row=0, column=1, sticky=tk.W)
        ttk.Label(opt, text="저장 폴더:").grid(row=1, column=0, padx=6, pady=6, sticky=tk.W)
        frm2 = ttk.Frame(opt)
        frm2.grid(row=1, column=1, sticky=tk.W)
        ttk.Entry(frm2, textvariable=self.out_var, width=40).pack(side=tk.LEFT)
        ttk.Button(frm2, text="찾아보기", command=lambda: self.out_var.set(filedialog.askdirectory() or self.out_var.get())).pack(side=tk.LEFT, padx=8)

        self.word_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(opt, text="보고서 생성 후 Word(.docx) 로도 저장", variable=self.word_var).grid(row=2, column=0, columnspan=2, sticky=tk.W, pady=6)
        self.chart_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(opt, text="📊 수치 데이터 자동 차트 생성 (LLM 추출 → Word 삽입)",
                        variable=self.chart_var).grid(row=3, column=0, columnspan=2, sticky=tk.W, pady=4)

        self.btn = ttk.Button(self, text="보고서 생성 (LLM)", command=self._generate)
        self.btn.grid(row=4, column=0, pady=12)
        self.prog = ttk.Progressbar(self, mode="indeterminate", length=700)
        self.prog.grid(row=5, column=0, columnspan=2, sticky=tk.EW)

        ttk.Label(self, text="로그:", style="Title.TLabel").grid(row=6, column=0, sticky=tk.W, pady=8)
        lf2 = ttk.Frame(self, style="Card.TFrame")
        lf2.grid(row=7, column=0, columnspan=2, sticky=tk.NSEW, pady=8)
        sb2 = ttk.Scrollbar(lf2)
        sb2.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_box = tk.Text(lf2, height=10, yscrollcommand=sb2.set, wrap=tk.WORD, bg=LOG_BG, fg=LOG_FG, font=("Consolas", 9))
        self.log_box.pack(fill=tk.BOTH, expand=True)
        sb2.config(command=self.log_box.yview)

    def log(self, msg):
        self.log_box.insert(tk.END, msg + "\n")
        self.log_box.see(tk.END)
        self.update_idletasks()

    # ── 파일 목록 관리 ────────────────────────────────────
    def _refresh_count(self):
        total   = len(self.md_files)
        checked = sum(1 for iid in self.file_list.get_children()
                      if self.file_list.item(iid, "values")[0] == "✅")
        self.file_count_lbl.config(text=f"총 {total}개  |  체크 {checked}개")

    def _insert_file(self, path):
        path = os.path.normpath(path)
        if path in self.md_files:
            return False
        self.md_files.append(path)
        fname  = os.path.basename(path)
        folder = os.path.basename(os.path.dirname(path)) or os.path.dirname(path)
        iid = self.file_list.insert("", tk.END,
                                    values=("✅", fname, folder),
                                    tags=("checked",))
        return True

    def _on_tree_click(self, event):
        """행 클릭 시 체크 토글"""
        iid = self.file_list.identify_row(event.y)
        if not iid:
            return
        vals = list(self.file_list.item(iid, "values"))
        if vals[0] == "✅":
            vals[0] = "☐"
            self.file_list.item(iid, values=vals, tags=("unchecked",))
        else:
            vals[0] = "✅"
            self.file_list.item(iid, values=vals, tags=("checked",))
        self._refresh_count()

    def _check_all(self):
        all_iids = self.file_list.get_children()
        for iid in all_iids:
            vals = list(self.file_list.item(iid, "values"))
            vals[0] = "✅"
            self.file_list.item(iid, values=vals, tags=("checked",))
        # Treeview 행도 하이라이트 → 선택 제거 버튼과 연동
        self.file_list.selection_set(all_iids)
        self._refresh_count()

    def _uncheck_all(self):
        all_iids = self.file_list.get_children()
        for iid in all_iids:
            vals = list(self.file_list.item(iid, "values"))
            vals[0] = "☐"
            self.file_list.item(iid, values=vals, tags=("unchecked",))
        self.file_list.selection_remove(all_iids)
        self._refresh_count()

    def _remove_selected(self):
        """Treeview 에서 하이라이트된 항목 제거"""
        sel = self.file_list.selection()
        if not sel:
            messagebox.showinfo("알림", "제거할 파일을 클릭하거나\n[전체 선택] 후 제거하세요.")
            return
        for iid in sel:
            vals  = self.file_list.item(iid, "values")
            fname = vals[1]
            folder_name = vals[2]
            # md_files 에서도 제거
            self.md_files = [p for p in self.md_files
                             if not (os.path.basename(p) == fname and
                                     os.path.basename(os.path.dirname(p)) == folder_name)]
            self.file_list.delete(iid)
        self.log(f"{len(sel)}개 파일을 제거했습니다.")
        self._refresh_count()

    def _clear_all(self):
        if not self.md_files:
            return
        if messagebox.askyesno("확인", "목록을 전부 비울까요?"):
            self.file_list.delete(*self.file_list.get_children())
            self.md_files.clear()
            self.log("목록을 초기화했습니다.")
            self._refresh_count()

    def _load_folder(self):
        folder = filedialog.askdirectory(title="MD 파일이 있는 폴더 선택")
        if not folder:
            return
        try:
            # 재귀 탐색 + 루트 직접 탐색 (Windows 경로 대응)
            files = set()
            files.update(glob.glob(os.path.join(folder, "*.md")))
            files.update(glob.glob(os.path.join(folder, "**", "*.md"), recursive=True))
            files = sorted(files)
            self.log(f"폴더 탐색: '{os.path.basename(folder)}' → {len(files)}개 MD 파일 발견")
            if not files:
                messagebox.showinfo("알림", f"폴더에 MD 파일이 없습니다.\n{folder}")
                return
            count = 0
            for f in files:
                try:
                    if self._insert_file(f):
                        count += 1
                except Exception as e:
                    self.log(f"  [오류] {os.path.basename(f)}: {e}")
            self.log(f"{count}개 파일 추가 완료.")
            self._refresh_count()
            self.update_idletasks()
        except Exception as e:
            self.log(f"[오류] 폴더 로드 실패: {e}")
            messagebox.showerror("오류", str(e))

    def _load_files(self):
        files = filedialog.askopenfilenames(
            title="MD 파일 선택",
            filetypes=[("Markdown files", "*.md"), ("All files", "*.*")])
        count = sum(1 for f in files if self._insert_file(f))
        self.log(f"{count}개 파일 추가.")
        self._refresh_count()

    def _get_active_files(self):
        """체크(✅)된 파일 경로 반환. 없으면 전체 반환."""
        checked = []
        for iid in self.file_list.get_children():
            vals = self.file_list.item(iid, "values")
            if vals[0] == "✅":
                fname  = vals[1]
                folder = vals[2]
                for p in self.md_files:
                    if (os.path.basename(p) == fname and
                            os.path.basename(os.path.dirname(p)) == folder):
                        checked.append(p)
                        break
        return checked if checked else list(self.md_files)

    def _generate(self):
        if not self.md_files:
            messagebox.showwarning("경고", "생성할 MD 파일이 없습니다.")
            return
        active = self._get_active_files()
        selected_contents = []
        for path in active:
            try:
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read()
                title = os.path.basename(path)
                if content.startswith("# "):
                    title = content.split("\n", 1)[0][2:].strip()
                selected_contents.append({"title": title, "content": content})
            except Exception as e:
                self.log(f"파일 읽기 실패: {path} - {e}")
        if not selected_contents:
            messagebox.showerror("오류", "모든 파일 읽기에 실패했습니다.")
            return
        self.btn.config(state="disabled")
        self.prog.start()
        checked_n = sum(1 for iid in self.file_list.get_children()
                         if self.file_list.item(iid, "values")[0] == "✅")
        hint = f"체크 {checked_n}개 / 전체 {len(self.md_files)}개"
        self.log(f"파일 종합 중... ({hint})")

        do_charts = self.chart_var.get()

        def _worker():
            try:
                report = llm_generate_report(selected_contents)
                charts = []
                if do_charts:
                    charts = llm_extract_chart_data(selected_contents, callback=self.log)
                self.after(0, lambda: self._save_report(report, chart_data=charts,
                                                         source_contents=selected_contents))
            except Exception as e:
                self.after(0, lambda: (self.log(f"[오류] {e}"), messagebox.showerror("오류", str(e)), self.prog.stop(), self.btn.config(state="normal")))

        threading.Thread(target=_worker, daemon=True).start()

    def _export_to_word(self):
        if not self.md_files:
            messagebox.showwarning("경고", "변환할 MD 파일이 없습니다.")
            return
        active = self._get_active_files()
        checked_n = sum(1 for iid in self.file_list.get_children()
                         if self.file_list.item(iid, "values")[0] == "✅")
        sel_hint = f"체크 {checked_n}개 / 전체 {len(self.md_files)}개"
        self.btn.config(state="disabled")
        self.prog.start()
        self.log(f"Word 변환 중 ({sel_hint})...")

        def _worker():
            try:
                out_dir = self.out_var.get().strip() or "./word_output"
                os.makedirs(out_dir, exist_ok=True)
                results = []
                for path in active:
                    try:
                        fname = os.path.basename(path).replace(".md", "")
                        out_path = os.path.join(out_dir, f"{fname}.docx")
                        result = generate_word_report(path, out_path,
                                                       title=fname, callback=self.log)
                        if result:
                            results.append(out_path)
                        else:
                            self.log(f"[Word] ⚠️ 변환 실패: {fname} (로그 확인)")
                    except Exception as e:
                        self.log(f"변환 실패: {path} - {e}")
                self.after(0, lambda: (self.prog.stop(), self.btn.config(state="normal"), self.log(f"Word 변환 완료: {len(results)}개"), messagebox.showinfo("완료", f"Word 변환 완료!\n{len(results)}개 파일 생성:\n{out_dir}")))
            except Exception as e:
                self.after(0, lambda: (self.log(f"[오류] {e}"), messagebox.showerror("오류", str(e)), self.prog.stop(), self.btn.config(state="normal")))

        threading.Thread(target=_worker, daemon=True).start()

    def _save_report(self, report: str, chart_data: list = None, source_contents: list = None):
        try:
            os.makedirs(self.out_var.get(), exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{self.report_title_var.get().replace(' ', '_')}_{timestamp}.md"
            path = os.path.join(self.out_var.get(), filename)
            # 차트 → MD 파일과 같은 폴더 내 임시 서브폴더에 생성 후 Word 삽입 후 삭제
            import shutil as _shutil2
            timestamp2 = datetime.now().strftime("%Y%m%d_%H%M%S%f")
            charts_tmp = os.path.join(os.path.abspath(self.out_var.get()), f"_charts_{timestamp2}") if chart_data else None
            chart_md_block = ""
            chart_paths = []
            try:
                if chart_data and charts_tmp:
                    chart_paths = generate_charts(chart_data, charts_tmp, callback=self.log)
                    if chart_paths:
                        chart_md_block = "\n\n## 📊 데이터 시각화\n\n"
                        for fig_no, (cp, spec) in enumerate(zip(chart_paths, chart_data), 1):
                            rel  = os.path.relpath(cp, os.path.abspath(self.out_var.get())).replace("\\", "/")
                            spec_title  = spec.get("title", f"차트 {fig_no}")
                            spec_src    = spec.get("source", "")
                            spec_ins    = spec.get("insight", "")
                            spec_labels = spec.get("labels", [])
                            spec_values = spec.get("values", [])
                            spec_unit   = spec.get("unit", "")
                            # 원본 데이터 표
                            if spec_labels and spec_values:
                                chart_md_block += "| 항목 | 수치 |\n|------|------|\n"
                                for lbl, val in zip(spec_labels, spec_values):
                                    chart_md_block += f"| {lbl} | {val}{spec_unit} |\n"
                                chart_md_block += "\n"
                            # 차트 이미지
                            chart_md_block += f"[EMBED_IMAGE:{rel}]\n\n"
                            # 캡션 + 출처
                            caption = f"**그림 {fig_no}. {spec_title}**"
                            if spec_src:
                                caption += f"  *(출처: {spec_src})*"
                            chart_md_block += caption + "\n\n"
                            # 인사이트 박스
                            if spec_ins:
                                chart_md_block += f"> 💡 **분석**: {spec_ins}\n\n"

                abs_path = os.path.abspath(path)
                with open(abs_path, "w", encoding="utf-8") as f:
                    f.write(f"# {self.report_title_var.get()}\n\n")
                    f.write(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                    f.write("---\n\n")
                    f.write(report)
                    if chart_md_block:
                        f.write(chart_md_block)
                if self.word_var.get():
                    out_docx = abs_path.replace(".md", ".docx")
                    word_ok = generate_word_report(abs_path, out_docx,
                                                      title=self.report_title_var.get(),
                                                      callback=self.log)
                    if not word_ok:
                        self.log("[Word] ⚠️ Word 생성 실패 — 로그를 확인하세요.")
            finally:
                # 차트 PNG는 Word에 삽입됐으므로 임시 폴더 삭제
                if charts_tmp and os.path.isdir(charts_tmp):
                    _shutil2.rmtree(charts_tmp, ignore_errors=True)
                    self.log("임시 차트 파일 정리 완료")
            self.prog.stop()
            self.btn.config(state="normal")
            self.log(f"보고서 생성 완료: {path}")
            messagebox.showinfo("완료", f"보고서 생성 완료!\n\n{path}")
        except Exception as e:
            self.prog.stop()
            self.btn.config(state="normal")
            self.log(f"[오류] 보고서 저장 실패: {e}")
            messagebox.showerror("오류", f"보고서 저장 실패:\n{e}")


class SettingsTab(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent, style="Card.TFrame", padding=20)
        self._build()

    def _build(self):
        self.columnconfigure(1, weight=1)
        ttk.Label(self, text="루코드 LLM 연결 설정", style="Title.TLabel").grid(row=0, column=0, columnspan=2, pady=(0, 16), sticky="w")

        ttk.Label(self, text="API Key:").grid(row=1, column=0, sticky="w", pady=8)
        key_frm = ttk.Frame(self, style="Card.TFrame")
        key_frm.grid(row=1, column=1, sticky="ew", pady=8)
        self.llm_key_var = tk.StringVar(value=LLM_API_KEY)
        self.key_entry = ttk.Entry(key_frm, textvariable=self.llm_key_var, show="*")
        self.key_entry.pack(side="left", fill="x", expand=True)
        self._show_key = False

        def toggle_key():
            self._show_key = not self._show_key
            self.key_entry.config(show="" if self._show_key else "*")
            show_btn.config(text="숨기기" if self._show_key else "보기")

        show_btn = ttk.Button(key_frm, text="보기", width=6, command=toggle_key)
        show_btn.pack(side="left", padx=6)

        ttk.Label(self, text="Base URL:").grid(row=2, column=0, sticky="w", pady=8)
        self.llm_url_var = tk.StringVar(value=LLM_BASE_URL)
        ttk.Entry(self, textvariable=self.llm_url_var).grid(row=2, column=1, sticky="ew", pady=8)
        ttk.Label(self, text="예) http://10.240.246.158:8000/v1", style="Muted.TLabel").grid(row=3, column=1, sticky="w")

        ttk.Label(self, text="모델명:").grid(row=4, column=0, sticky="w", pady=8)
        self.llm_model_var = tk.StringVar(value=LLM_MODEL)
        ttk.Entry(self, textvariable=self.llm_model_var, width=32).grid(row=4, column=1, sticky="w", pady=8)

        ttk.Label(self, text="Vision/OCR 후처리 모델:").grid(row=5, column=0, sticky="w", pady=8)
        self.llm_vision_var = tk.StringVar(value=LLM_VISION_MODEL)
        ttk.Entry(self, textvariable=self.llm_vision_var, width=32).grid(row=5, column=1, sticky="w", pady=8)
        ttk.Label(self, text="현재 코드는 이미지를 직접 보내지 않고 OCR 텍스트를 Qwen에 보냅니다.", style="Muted.TLabel").grid(row=6, column=1, sticky="w")

        ttk.Label(self, text="최대 토큰:").grid(row=7, column=0, sticky="w", pady=8)
        self.llm_tokens_var = tk.StringVar(value=str(LLM_MAX_TOKENS))
        ttk.Entry(self, textvariable=self.llm_tokens_var, width=10).grid(row=7, column=1, sticky="w", pady=8)

        btn_frm = ttk.Frame(self, style="Card.TFrame")
        btn_frm.grid(row=8, column=0, columnspan=2, pady=16, sticky="w")
        ttk.Button(btn_frm, text="설정 저장", command=self._save_settings).pack(side="left", padx=6)
        ttk.Button(btn_frm, text="연결 테스트", command=self._test_llm).pack(side="left", padx=6)

        self.settings_status = ttk.Label(self, text="", style="Muted.TLabel")
        self.settings_status.grid(row=9, column=0, columnspan=2, sticky="w", pady=8)

        ttk.Separator(self, orient="horizontal").grid(row=10, column=0, columnspan=2, sticky="ew", pady=12)
        info = (
            "Vision API 이미지 분석 안내\n"
            "이미지를 base64로 인코딩하여 Qwen Vision 모델에 직접 전달합니다.\n"
            "별도 OCR 라이브러리 설치 없이 동작합니다.\n"
            "Vision 모델 설정: LLM 설정 탭의 Vision 모델 항목을 확인하세요."
        )
        ttk.Label(self, text=info, style="Muted.TLabel", justify="left").grid(row=11, column=0, columnspan=2, sticky="w")

    def _save_settings(self):
        key = self.llm_key_var.get().strip()
        url = self.llm_url_var.get().strip()
        model = self.llm_model_var.get().strip()
        vision = self.llm_vision_var.get().strip()
        tokens = self.llm_tokens_var.get().strip()
        if not url:
            self.settings_status.config(text="Base URL 은 필수입니다.", foreground="red")
            return
        try:
            save_llm_config(key, url, model, vision, int(tokens))
            self.settings_status.config(text=f"저장 완료 -> {CONFIG_FILE}", foreground="green")
            messagebox.showinfo("완료", "설정이 저장되었습니다.")
        except Exception as e:
            self.settings_status.config(text=f"저장 실패: {e}", foreground="red")

    def _test_llm(self):
        self.settings_status.config(text="연결 테스트 중...", foreground="blue")
        self.update_idletasks()
        key = self.llm_key_var.get().strip()
        url = self.llm_url_var.get().strip()
        model = self.llm_model_var.get().strip()

        def _test():
            try:
                OpenAI = _need_openai_class()
                client = OpenAI(api_key=key if key else "sk-ignored", base_url=url)
                resp = client.chat.completions.create(
                    model=model,
                    messages=[{"role": "user", "content": "연결 테스트입니다. 한 문장으로 답해주세요."}],
                    max_tokens=100,
                    temperature=0.1,
                )
                answer = resp.choices[0].message.content.strip()[:120]
                self.after(0, lambda: self.settings_status.config(text=f"연결 성공! 응답: {answer}", foreground="green"))
            except Exception as e:
                self.after(0, lambda: self.settings_status.config(text=f"연결 실패: {type(e).__name__}: {e}", foreground="red"))

        threading.Thread(target=_test, daemon=True).start()


if __name__ == "__main__":
    App().mainloop()
