"""
Confluence 수집기 GUI - v6
신규: Qwen Vision 이미지 분석 + Word(.docx) 보고서 + 자동 차트 생성
"""
import subprocess, sys

def _install(pkg):
    subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q"],
                          stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

for _pkg, _imp in [("html2text","html2text"), ("openai","openai"),
                   ("python-docx","docx"), ("matplotlib","matplotlib"),
                   ("Pillow","PIL"), ("beautifulsoup4","bs4")]:
    try:
        __import__(_imp)
    except ImportError:
        print(f"{_pkg} 설치 중..."); _install(_pkg)

import os, re, threading, traceback, json, base64, tempfile
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from typing import Dict, List, Optional
from datetime import datetime
from bs4 import BeautifulSoup
from openai import OpenAI
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 색상 팔레트 ───────────────────────────────
BLUE    = "#1428A0"; BLUE_DK = "#0D1E7A"; BLUE_LT = "#E8EAFF"
CYAN    = "#00B0FF"; BG      = "#F0F2F7"; CARD    = "#FFFFFF"
BORDER  = "#D0D5E0"; TEXT    = "#1A1A2E"; TEXT_MUTED = "#6B7280"
GREEN   = "#16A34A"; RED     = "#DC2626"; ORANGE  = "#EA580C"
LOG_BG  = "#0D1117"; LOG_FG  = "#C9D1D9"
LOG_GREEN="#3FB950"; LOG_RED ="#FF7B72"
LOG_BLUE="#79C0FF";  LOG_YELLOW="#E3B341"

# ─── 설정 ──────────────────────────────────────
BASE_URL      = "https://confluence.sec.samsung.net"
USER_DATA_DIR = "./chrome_profile_confluence_md"
CONFIG_FILE   = os.path.join(os.path.dirname(os.path.abspath(__file__)), "llm_config.json")
LLM_API_KEY = LLM_BASE_URL = ""; LLM_MODEL = "qwen-plus"; LLM_MAX_TOKENS = 2000

def load_llm_config():
    global LLM_API_KEY, LLM_BASE_URL, LLM_MODEL, LLM_MAX_TOKENS
    if os.path.exists(CONFIG_FILE):
        try:
            c = json.loads(open(CONFIG_FILE, encoding="utf-8").read())
            LLM_API_KEY=c.get("api_key",""); LLM_BASE_URL=c.get("base_url","")
            LLM_MODEL=c.get("model","qwen-plus"); LLM_MAX_TOKENS=int(c.get("max_tokens",2000))
        except: pass

def save_llm_config(k,u,m,t):
    global LLM_API_KEY,LLM_BASE_URL,LLM_MODEL,LLM_MAX_TOKENS
    LLM_API_KEY=k; LLM_BASE_URL=u; LLM_MODEL=m; LLM_MAX_TOKENS=int(t)
    with open(CONFIG_FILE,"w",encoding="utf-8") as f:
        json.dump({"api_key":k,"base_url":u,"model":m,"max_tokens":t},f,ensure_ascii=False,indent=2)

load_llm_config()
def _llm_ready(): return bool(LLM_API_KEY and LLM_BASE_URL)

# ─── LLM 함수 ──────────────────────────────────
def _client(): return OpenAI(api_key=LLM_API_KEY, base_url=LLM_BASE_URL)

def llm_summarize_page(title: str, text: str) -> str:
    try:
        resp = _client().chat.completions.create(
            model=LLM_MODEL, temperature=0.1, max_tokens=LLM_MAX_TOKENS,
            messages=[
                {"role":"system","content":
                    "당신은 400명 규모 조직의 팀장을 위한 업무 보고서 작성 전문가입니다.\n"
                    "【절대 규칙】제공된 문서 내용만 사용. 없는 내용 절대 금지.\n"
                    "【작성】전문 용어는 괄호로 설명. 격식체. 수치/날짜 그대로.\n"
                    "【형식】## 개요 / ## 주요 내용 / ## 완료 / ## 진행 중 / ## 이슈"},
                {"role":"user","content":f"제목: {title}\n\n{text[:5000]}"}])
        return resp.choices[0].message.content
    except Exception as e: return f"[요약 실패: {e}]"

def llm_analyze_image(image_b64: str, mime: str = "image/png", context: str = "") -> str:
    """Qwen Vision으로 이미지 분석"""
    try:
        data_url = f"data:{mime};base64,{image_b64}"
        resp = _client().chat.completions.create(
            model=LLM_MODEL, temperature=0.1, max_tokens=500,
            messages=[{"role":"user","content":[
                {"type":"text","text":
                    f"이 이미지를 업무 보고서 관점에서 분석해주세요.\n"
                    f"{'페이지 맥락: ' + context if context else ''}\n"
                    "- 이미지에 포함된 핵심 정보 (수치, 상태, 내용) 설명\n"
                    "- 차트/그래프면 수치 읽기\n"
                    "- 다이어그램이면 구조/흐름 설명\n"
                    "- 3~5문장, 격식체"},
                {"type":"image_url","image_url":{"url": data_url}}
            ]}])
        return resp.choices[0].message.content
    except Exception as e: return f"[이미지 분석 실패: {e}]"

def llm_generate_report(pages: list) -> str:
    combined = "".join(f"\n\n### {p['title']}\n{p['content'][:2000]}" for p in pages)
    try:
        resp = _client().chat.completions.create(
            model=LLM_MODEL, temperature=0.1, max_tokens=LLM_MAX_TOKENS*2,
            messages=[
                {"role":"system","content":
                    "당신은 400명 규모 조직의 팀장을 위한 종합 보고서 작성 전문가입니다.\n"
                    "제공된 내용만 사용. 추측/창작 금지.\n"
                    "전문 용어 괄호 설명. 격식 문어체. 수치 반드시 포함.\n"
                    "# 1. 전체 요약\n# 2. 항목별 상세 현황\n"
                    "# 3. 완료된 주요 사항\n# 4. 진행 중인 과제\n"
                    "# 5. 이슈 및 리스크\n# 6. 팀장 조치 필요 사항"},
                {"role":"user","content":
                    f"{len(pages)}개 페이지로 팀장 보고용 종합 보고서를 작성해주세요.\n\n{combined}"}])
        return resp.choices[0].message.content
    except Exception as e: return f"[보고서 실패: {e}]"

def llm_extract_chart_data(report_text: str) -> dict:
    """보고서에서 차트용 수치 데이터 추출"""
    try:
        resp = _client().chat.completions.create(
            model=LLM_MODEL, temperature=0, max_tokens=600,
            messages=[
                {"role":"system","content":
                    "보고서 텍스트에서 차트로 그릴 수 있는 수치 데이터를 JSON으로 추출하세요.\n"
                    "형식: {\"charts\": [{\"title\":\"차트제목\", \"type\":\"bar|pie\", "
                    "\"labels\":[\"항목1\",\"항목2\"], \"values\":[숫자1,숫자2]}]}\n"
                    "차트가 없으면 {\"charts\":[]} 반환. 반드시 JSON만 반환."},
                {"role":"user","content": report_text[:3000]}])
        raw = resp.choices[0].message.content.strip()
        raw = re.sub(r"```json|```","",raw).strip()
        return json.loads(raw)
    except: return {"charts":[]}

# ─── 유틸 ──────────────────────────────────────
def clean_filename(t): return re.sub(r'[\\/*?:"<>|]',"_",t).strip()

def extract_page_id_from_url(url):
    for p in [r'/pages/(\d+)',r'pageId=(\d+)',r'pages/(\d{8,})']:
        m = re.search(p,url)
        if m: return m.group(1)
    return None

def analyze_confluence_content(html: str) -> Dict:
    soup = BeautifulSoup(html,'html.parser')
    items=[]
    for i in range(1,7):
        for h in soup.find_all(f'h{i}'):
            items.append({'type':f'heading_{i}','content':h.get_text(strip=True)})
    for table in soup.find_all('table'):
        headers=[th.get_text(strip=True) for th in table.find_all('th')]
        rows=[[td.get_text(strip=True) for td in tr.find_all('td')]
              for tr in table.find_all('tr') if tr.find_all('td')]
        if headers or rows: items.append({'type':'table','headers':headers,'rows':rows})
    for ul in soup.find_all(['ul','ol']):
        li=[l.get_text(strip=True) for l in ul.find_all('li',recursive=False)]
        if li: items.append({'type':'list','items':li})
    for p in soup.find_all('p'):
        t=p.get_text(strip=True)
        if t and len(t)>10: items.append({'type':'paragraph','content':t})
    images=[{'src':img.get('src',''),'alt':img.get('alt','')}
            for img in soup.find_all('img') if img.get('src')]
    return {'text_content':items,'images':images}

def convert_to_markdown(analysis: Dict, title: str, image_analyses: list=None) -> str:
    md=[f"# {title}","",f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}","---",""]
    if image_analyses:
        md+=["## 이미지 분석",""]
        for idx,(img,analysis_text) in enumerate(image_analyses,1):
            md.append(f"### 이미지 {idx}: {img.get('alt','') or '(이미지)'}")
            md.append(analysis_text)
            md.append("")
    md+=["## 내용",""]
    for item in analysis['text_content']:
        t=item['type']
        if t.startswith('heading_'):
            md+=[f"{'#'*int(t[-1])} {item['content']}",""]
        elif t=='table':
            h=item.get('headers',[])
            if h:
                md+=["| "+" | ".join(h)+" |","| "+" | ".join(["---"]*len(h))+" |"]
                for row in item.get('rows',[]):
                    while len(row)<len(h): row.append("")
                    md.append("| "+" | ".join(row)+" |")
                md.append("")
        elif t=='list':
            for it in item.get('items',[]): md.append(f"- {it}")
            md.append("")
        elif t=='paragraph':
            md+=[item['content'],""]
    return "\n".join(md)

def get_page_children(page, pid):
    all_p,start=[],0
    while True:
        r=page.request.get(f"{BASE_URL}/rest/api/content/search",
            params={"cql":f"ancestor={pid} and type=page","start":str(start),"limit":"50","expand":"version"})
        if r.status!=200: break
        data=r.json(); results=data.get("results",[])
        if not results: break
        all_p+=[{"id":d["id"],"title":d["title"]} for d in results]
        if len(all_p)>=data.get("size",0): break
        start+=50
        if start>500: break
    return all_p

def get_page_content(page, pid):
    try:
        r=page.request.get(f"{BASE_URL}/rest/api/content/{pid}",params={"expand":"body.storage"})
        if r.status!=200: return None
        d=r.json(); html=d.get("body",{}).get("storage",{}).get("value","")
        return {"id":d.get("id"),"title":d.get("title"),"html":html} if html else None
    except: return None

def download_image_b64(page, url: str) -> tuple:
    """Playwright 세션으로 이미지 다운로드 → (base64, mime)"""
    try:
        if url.startswith("/"): url = BASE_URL + url
        r = page.request.get(url)
        if r.status != 200: return None, None
        body = r.body()
        ct = r.headers.get("content-type","image/png").split(";")[0].strip()
        return base64.b64encode(body).decode(), ct
    except: return None, None

def process_child_pages_recursive(page, page_id, save_dir,
                                   depth=2, use_llm=True, use_vision=True, callback=None):
    if callback: callback(f"INFO: 페이지 {page_id} 처리 중...")
    try:
        data = get_page_content(page, page_id)
        if not data:
            if callback: callback(f"WARN: {page_id} 내용 없음"); return
        analysis = analyze_confluence_content(data['html'])

        # 이미지 Vision 분석
        image_analyses = []
        if use_vision and use_llm and _llm_ready() and analysis['images']:
            if callback: callback(f"INFO: 이미지 {len(analysis['images'])}개 Vision 분석 중...")
            for img in analysis['images'][:5]:  # 최대 5개
                b64, mime = download_image_b64(page, img['src'])
                if b64:
                    if callback: callback(f"INFO:   - {img.get('alt','이미지')} 분석 중...")
                    desc = llm_analyze_image(b64, mime, context=data['title'])
                    image_analyses.append((img, desc))

        markdown = convert_to_markdown(analysis, data['title'], image_analyses or None)
        llm_block = ""
        if use_llm and _llm_ready():
            if callback: callback(f"INFO: LLM 요약 중: {data['title']}")
            summary = llm_summarize_page(data['title'], markdown)
            llm_block = f"\n\n---\n## AI 요약 (팀장 보고용)\n\n{summary}\n\n---\n"

        safe = clean_filename(data['title'])
        path = os.path.join(save_dir, f"{page_id}_{safe}.md")
        with open(path,"w",encoding="utf-8") as f: f.write(markdown+llm_block)
        if callback: callback(f"OK: {data['title']}")

        children = get_page_children(page, page_id)
        if children and depth>0:
            cdir = os.path.join(save_dir,f"children_{page_id}")
            os.makedirs(cdir,exist_ok=True)
            for child in children:
                process_child_pages_recursive(page,child['id'],cdir,depth-1,use_llm,use_vision,callback)
    except Exception as e:
        if callback: callback(f"ERR: {e}")

def generate_report_from_md_files(md_dir, use_llm=True, callback=None) -> str:
    pages=[]
    for root,_,files in os.walk(md_dir):
        for fname in sorted(files):
            if fname.endswith(".md") and "보고서" not in fname and "report" not in fname.lower():
                path=os.path.join(root,fname)
                with open(path,encoding="utf-8") as f: content=f.read()
                pages.append({"title":fname.replace(".md",""),"content":content,"path":path})
    if callback: callback(f"INFO: MD 파일 {len(pages)}개 로드")
    if not pages: return "보고서 생성 실패: MD 파일 없음"
    if use_llm and _llm_ready():
        if callback: callback("INFO: LLM 종합 보고서 생성 중...")
        return llm_generate_report(pages)
    lines=["# 통합 보고서\n",f"총 {len(pages)}개 문서\n"]
    for p in pages:
        lines+=[f"\n## {p['title']}\n",p['content'],"\n---\n"]
    return "\n".join(lines)

# ─── Word 보고서 생성 ───────────────────────────
def _make_chart(chart_info: dict, save_path: str) -> bool:
    try:
        labels=chart_info.get("labels",[]); values=chart_info.get("values",[])
        title=chart_info.get("title",""); ctype=chart_info.get("type","bar")
        if not labels or not values: return False

        # 한글 폰트 설정
        font_candidates = ["Malgun Gothic","AppleGothic","NanumGothic","DejaVu Sans"]
        for fc in font_candidates:
            if any(fc.lower() in f.name.lower() for f in fm.fontManager.ttflist):
                plt.rcParams['font.family'] = fc; break

        fig, ax = plt.subplots(figsize=(6,3.5), dpi=120)
        colors = ["#1428A0","#00B0FF","#16A34A","#EA580C","#9333EA","#DC2626"]
        if ctype == "pie":
            ax.pie(values, labels=labels, autopct='%1.1f%%',
                   colors=colors[:len(values)], startangle=90)
        else:
            bars = ax.bar(labels, values, color=colors[:len(labels)], width=0.55)
            for bar, val in zip(bars, values):
                ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.3,
                        str(val), ha='center', va='bottom', fontsize=9)
            ax.set_ylim(0, max(values)*1.2 if values else 10)
            ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
        ax.set_title(title, fontsize=11, pad=10)
        plt.tight_layout()
        plt.savefig(save_path, bbox_inches='tight', facecolor='white')
        plt.close(); return True
    except: return False

def create_word_report(report_text: str, title: str, output_path: str,
                       use_charts: bool = True, callback=None) -> str:
    doc = Document()
    for section in doc.sections:
        section.page_width=Cm(21); section.page_height=Cm(29.7)
        section.left_margin=section.right_margin=Cm(2.5)
        section.top_margin=section.bottom_margin=Cm(2.5)

    # 헤더 스타일 함수
    def add_h(text, level=1):
        p = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.name="맑은 고딕"
        if level==1:
            run.font.size=Pt(18); run.font.color.rgb=RGBColor(0x14,0x28,0xA0)
        elif level==2:
            run.font.size=Pt(14); run.font.color.rgb=RGBColor(0x1F,0x49,0x7D)
        else:
            run.font.size=Pt(12); run.font.color.rgb=RGBColor(0x2E,0x74,0xB5)

    def add_p(text, bold=False, size=11):
        p=doc.add_paragraph()
        run=p.add_run(text)
        run.font.name="맑은 고딕"; run.font.size=Pt(size); run.bold=bold
        return p

    def add_bullet(text):
        p=doc.add_paragraph(style="List Bullet")
        run=p.add_run(text)
        run.font.name="맑은 고딕"; run.font.size=Pt(11)

    def add_md_table(headers, rows):
        if not headers: return
        table=doc.add_table(rows=1+len(rows), cols=len(headers))
        table.style="Table Grid"
        hrow=table.rows[0]
        for i,h in enumerate(headers):
            cell=hrow.cells[i]; cell.text=h
            run=cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
            run.bold=True; run.font.name="맑은 고딕"; run.font.size=Pt(10)
            cell._tc.get_or_add_tcPr()
        for ri,row in enumerate(rows):
            for ci,val in enumerate(row[:len(headers)]):
                cell=table.rows[ri+1].cells[ci]; cell.text=val
                for run in cell.paragraphs[0].runs:
                    run.font.name="맑은 고딕"; run.font.size=Pt(10)
        doc.add_paragraph()

    # 제목 + 메타
    add_h(title, 1)
    add_p(f"작성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}",size=9)
    doc.add_paragraph()

    # 차트 생성
    chart_images = []
    if use_charts and _llm_ready():
        if callback: callback("INFO: 차트 데이터 추출 중...")
        chart_data = llm_extract_chart_data(report_text)
        tmpdir = tempfile.mkdtemp()
        for i,c in enumerate(chart_data.get("charts",[])):
            p = os.path.join(tmpdir,f"chart_{i}.png")
            if _make_chart(c,p):
                chart_images.append((c.get("title","차트"), p))
                if callback: callback(f"OK: 차트 생성 '{c.get('title','')}'")

    # 보고서 텍스트 파싱 → Word
    current_table_headers = []
    current_table_rows = []

    def flush_table():
        if current_table_headers or current_table_rows:
            add_md_table(current_table_headers, current_table_rows)
            current_table_headers.clear(); current_table_rows.clear()

    for line in report_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            flush_table(); doc.add_paragraph(); continue
        if stripped.startswith("# "):
            flush_table(); add_h(stripped[2:], 1)
        elif stripped.startswith("## "):
            flush_table(); add_h(stripped[3:], 2)
        elif stripped.startswith("### "):
            flush_table(); add_h(stripped[4:], 3)
        elif stripped.startswith(("- ","• ","* ")):
            flush_table(); add_bullet(stripped[2:])
        elif re.match(r"^\d+\.", stripped):
            flush_table()
            p=doc.add_paragraph(style="List Number")
            run=p.add_run(stripped.split(".",1)[-1].strip())
            run.font.name="맑은 고딕"; run.font.size=Pt(11)
        elif stripped.startswith("|"):
            # MD 테이블 파싱
            cells=[c.strip() for c in stripped.strip("|").split("|")]
            if all(re.match(r"^-+$",c) for c in cells if c): continue  # 구분선 스킵
            if not current_table_headers:
                current_table_headers = cells
            else:
                current_table_rows.append(cells)
        else:
            flush_table(); add_p(stripped)

    flush_table()

    # 차트 삽입
    if chart_images:
        doc.add_page_break()
        add_h("📊 데이터 차트", 1)
        for chart_title, chart_path in chart_images:
            add_p(chart_title, bold=True)
            try: doc.add_picture(chart_path, width=Inches(5.5))
            except: pass
            doc.add_paragraph()

    # 푸터
    for section in doc.sections:
        footer=section.footer
        fp=footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text=f"자동 생성 보고서  |  {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for run in fp.runs:
            run.font.size=Pt(8); run.font.color.rgb=RGBColor(0x70,0x70,0x70)

    doc.save(output_path)
    # 임시 차트 파일 정리
    for _,p in chart_images:
        try: os.remove(p)
        except: pass
    return output_path


# ─── 커스텀 버튼 ───────────────────────────────
class FlatButton(tk.Canvas):
    def __init__(self, parent, text, command=None, bg=BLUE, fg="white",
                 hover=BLUE_DK, width=140, height=36, font_size=10, **kw):
        super().__init__(parent,width=width,height=height,
                         highlightthickness=0,bd=0,cursor="hand2",**kw)
        self._bg=bg;self._hover=hover;self._fg=fg
        self._text=text;self._cmd=command;self._font=("Malgun Gothic",font_size,"bold")
        self._draw(bg)
        self.bind("<Enter>",lambda e:self._draw(hover))
        self.bind("<Leave>",lambda e:self._draw(bg))
        self.bind("<Button-1>",lambda e:self._click())

    def _draw(self,color):
        self.delete("all")
        w,h=int(self["width"]),int(self["height"]);r=6
        for x,y,s,e in [(0,0,90,90),(w-r*2,0,0,90),(0,h-r*2,180,90),(w-r*2,h-r*2,270,90)]:
            self.create_arc(x,y,x+r*2,y+r*2,start=s,extent=e,fill=color,outline=color)
        self.create_rectangle(r,0,w-r,h,fill=color,outline=color)
        self.create_rectangle(0,r,w,h-r,fill=color,outline=color)
        self.create_text(w//2,h//2,text=self._text,fill=self._fg,font=self._font)

    def _click(self):
        self._draw(self._bg)
        if self._cmd: self._cmd()

    def config_state(self,state):
        if state=="disabled":
            self._bg_orig=self._bg;self._bg="#9CA3AF"
            self._draw("#9CA3AF");self.unbind("<Button-1>")
        else:
            self._bg=getattr(self,'_bg_orig',BLUE);self._draw(self._bg)
            self.bind("<Button-1>",lambda e:self._click())


# ─── 메인 GUI ──────────────────────────────────
class ConfluenceGUI:
    def __init__(self, root):
        self.root=root
        self.root.title("Confluence 수집기  v6")
        self.root.geometry("1060x780"); self.root.configure(bg=BG)
        self.root.resizable(True,True)

        self.url_var      =tk.StringVar(); self.depth_var=tk.IntVar(value=7)
        self.out_var      =tk.StringVar(value="./confluence_output")
        self.llm_var      =tk.BooleanVar(value=True)
        self.vision_var   =tk.BooleanVar(value=True)
        self.word_var     =tk.BooleanVar(value=True)
        self.chart_var    =tk.BooleanVar(value=True)
        self.llm_key_var  =tk.StringVar(value=LLM_API_KEY)
        self.llm_url_var  =tk.StringVar(value=LLM_BASE_URL)
        self.llm_model_var=tk.StringVar(value=LLM_MODEL)
        self.llm_tok_var  =tk.StringVar(value=str(LLM_MAX_TOKENS))
        self._llm_open=False

        self._build()

    def _build(self):
        self._build_header()
        body=tk.Frame(self.root,bg=BG)
        body.pack(fill="both",expand=True,padx=14,pady=(10,0))
        body.columnconfigure(0,weight=0,minsize=360)
        body.columnconfigure(1,weight=1); body.rowconfigure(0,weight=1)
        self._build_settings(body); self._build_log(body)
        self._build_statusbar()

    def _build_header(self):
        hdr=tk.Frame(self.root,bg=BLUE,height=58); hdr.pack(fill="x"); hdr.pack_propagate(False)
        tk.Label(hdr,text="  Confluence 수집 & 보고서 생성",bg=BLUE,fg="white",
                 font=("Malgun Gothic",14,"bold")).pack(side="left",padx=16)
        tk.Label(hdr,text=" v6 ",bg=CYAN,fg=BLUE_DK,
                 font=("Malgun Gothic",9,"bold"),padx=4).pack(side="left",pady=16)
        tk.Label(hdr,text=" Vision + Word + 차트 ",bg="#1E3A8A",fg="#93C5FD",
                 font=("Malgun Gothic",8),padx=6).pack(side="left",pady=18)
        self._hdr_status=tk.Label(hdr,text="●  준비됨",bg=BLUE,fg="#86EFAC",
                                  font=("Malgun Gothic",9))
        self._hdr_status.pack(side="right",padx=20)

    def _build_settings(self,parent):
        outer=tk.Frame(parent,bg=CARD,highlightthickness=1,highlightbackground=BORDER)
        outer.grid(row=0,column=0,sticky="nsew",padx=(0,10),pady=(0,10))
        tk.Label(outer,text="⚙  설정",bg=CARD,fg=BLUE,
                 font=("Malgun Gothic",10,"bold")).pack(anchor="w",padx=16,pady=(12,0))
        tk.Frame(outer,bg=BLUE_LT,height=1).pack(fill="x",padx=16,pady=4)
        f=tk.Frame(outer,bg=CARD); f.pack(fill="x",padx=16)

        self._lbl(f,"Confluence URL")
        self.url_entry=self._entry(f,self.url_var); self.url_entry.pack(fill="x",pady=(0,2))
        self.url_var.trace_add("write",self._on_url)
        self.pid_lbl=tk.Label(f,text="페이지 ID: —",bg=CARD,fg=TEXT_MUTED,
                              font=("Malgun Gothic",8)); self.pid_lbl.pack(anchor="w",pady=(0,6))

        self._lbl(f,"재귀 깊이")
        tk.Spinbox(f,from_=0,to=15,textvariable=self.depth_var,width=5,
                   font=("Malgun Gothic",10),bd=1,relief="solid").pack(anchor="w",pady=(0,6))

        self._lbl(f,"출력 폴더")
        ofr=tk.Frame(f,bg=CARD); ofr.pack(fill="x",pady=(0,6))
        self._entry(ofr,self.out_var).pack(side="left",fill="x",expand=True)
        tk.Button(ofr,text="찾기",command=self._browse,bg=BG,fg=TEXT,
                  relief="flat",cursor="hand2",font=("Malgun Gothic",9),padx=6).pack(side="left",padx=4)

        # 옵션 체크박스들
        tk.Frame(f,bg=BORDER,height=1).pack(fill="x",pady=6)
        opts=[
            (self.llm_var,   "🤖 LLM 요약 생성"),
            (self.vision_var,"👁 이미지 Vision 분석 (Qwen)"),
            (self.word_var,  "📄 Word(.docx) 보고서 출력"),
            (self.chart_var, "📊 자동 차트 생성"),
        ]
        for var,lbl in opts:
            tk.Checkbutton(f,text=lbl,variable=var,bg=CARD,fg=TEXT,
                           activebackground=CARD,font=("Malgun Gothic",9),
                           selectcolor=BLUE_LT).pack(anchor="w",pady=1)

        # LLM 설정 패널
        tk.Frame(f,bg=BORDER,height=1).pack(fill="x",pady=8)
        self._llm_tog=tk.Label(f,text="▶  LLM 연결 설정",bg=BLUE_LT,fg=BLUE,
                               font=("Malgun Gothic",9,"bold"),cursor="hand2",padx=8,pady=4,anchor="w")
        self._llm_tog.pack(fill="x")
        self._llm_tog.bind("<Button-1>",lambda e:self._toggle_llm())
        self._llm_panel=tk.Frame(f,bg="#F5F7FF",highlightthickness=1,highlightbackground=BLUE_LT)
        self._build_llm_panel(self._llm_panel)

        # 진행바 + 버튼
        tk.Frame(outer,bg=BORDER,height=1).pack(fill="x",padx=16,pady=6)
        self.progress=ttk.Progressbar(outer,mode="indeterminate",length=320)
        self.progress.pack(fill="x",padx=16,pady=(0,6))
        bf=tk.Frame(outer,bg=CARD); bf.pack(pady=(0,16),padx=16,fill="x")
        self.run_btn=FlatButton(bf,text="▶  수집 시작",command=self._run,
                                bg=BLUE,fg="white",width=160,height=38,font_size=10)
        self.run_btn.pack(side="left")

    def _build_llm_panel(self,f):
        for lbl,var,show in [
            ("API Key",self.llm_key_var,"*"),
            ("Base URL",self.llm_url_var,""),
            ("모델명",self.llm_model_var,""),
            ("최대 토큰",self.llm_tok_var,""),
        ]:
            tk.Label(f,text=lbl,bg="#F5F7FF",fg=TEXT_MUTED,
                     font=("Malgun Gothic",8)).pack(anchor="w",padx=8,pady=(4,0))
            e=tk.Entry(f,textvariable=var,show=show,font=("Malgun Gothic",9),
                       bd=1,relief="solid",bg="white"); e.pack(fill="x",padx=8,pady=(0,2))
        br=tk.Frame(f,bg="#F5F7FF"); br.pack(fill="x",padx=8,pady=6)
        FlatButton(br,text="💾 저장",command=self._save_llm,
                   bg=GREEN,fg="white",width=80,height=30,font_size=9).pack(side="left")
        FlatButton(br,text="🔗 테스트",command=self._test_llm,
                   bg=BLUE,fg="white",width=90,height=30,font_size=9).pack(side="left",padx=6)
        self._llm_st=tk.Label(f,text="",bg="#F5F7FF",font=("Malgun Gothic",8),
                              wraplength=300,justify="left")
        self._llm_st.pack(anchor="w",padx=8,pady=(0,6))

    def _build_log(self,parent):
        lo=tk.Frame(parent,bg=LOG_BG,highlightthickness=1,highlightbackground="#30363D")
        lo.grid(row=0,column=1,sticky="nsew",pady=(0,10))
        hdr=tk.Frame(lo,bg="#161B22"); hdr.pack(fill="x")
        tk.Label(hdr,text="  로그",bg="#161B22",fg=LOG_BLUE,
                 font=("Consolas",9,"bold")).pack(side="left",pady=6,padx=8)
        tk.Button(hdr,text="지우기",command=self._clear_log,
                  bg="#161B22",fg="#6E7681",relief="flat",cursor="hand2",
                  font=("Consolas",8),padx=6).pack(side="right",padx=8)
        tf=tk.Frame(lo,bg=LOG_BG); tf.pack(fill="both",expand=True,padx=4,pady=4)
        sb=tk.Scrollbar(tf,bg=LOG_BG,troughcolor=LOG_BG); sb.pack(side="right",fill="y")
        self.log_box=tk.Text(tf,bg=LOG_BG,fg=LOG_FG,font=("Consolas",9),
                             wrap="word",yscrollcommand=sb.set,
                             insertbackground=LOG_FG,bd=0,relief="flat",
                             selectbackground="#264F78")
        self.log_box.pack(fill="both",expand=True); sb.config(command=self.log_box.yview)
        self.log_box.tag_config("ok",foreground=LOG_GREEN)
        self.log_box.tag_config("err",foreground=LOG_RED)
        self.log_box.tag_config("info",foreground=LOG_BLUE)
        self.log_box.tag_config("warn",foreground=LOG_YELLOW)
        self.log_box.tag_config("ts",foreground="#484F58")
        self.log_box.tag_config("done",foreground=LOG_GREEN,font=("Consolas",9,"bold"))

    def _build_statusbar(self):
        bar=tk.Frame(self.root,bg="#E5E7EB",height=24); bar.pack(fill="x",side="bottom")
        bar.pack_propagate(False)
        self._st_lbl=tk.Label(bar,text="  준비됨",bg="#E5E7EB",fg=TEXT_MUTED,
                              font=("Malgun Gothic",8),anchor="w")
        self._st_lbl.pack(side="left",fill="both",expand=True)
        tk.Label(bar,text="  Confluence Collector v6  Vision+Word+Chart  ",
                 bg="#E5E7EB",fg=TEXT_MUTED,font=("Malgun Gothic",8)).pack(side="right")

    # ─── 헬퍼 ──────────────────────────────────
    def _lbl(self,p,t):
        tk.Label(p,text=t,bg=CARD,fg=TEXT_MUTED,font=("Malgun Gothic",8)).pack(anchor="w",pady=(4,1))

    def _entry(self,p,v,width=30):
        return tk.Entry(p,textvariable=v,width=width,font=("Malgun Gothic",9),
                        bd=1,relief="solid",bg="white",fg=TEXT)

    def _toggle_llm(self):
        self._llm_open=not self._llm_open
        if self._llm_open:
            self._llm_panel.pack(fill="x",pady=(2,0))
            self._llm_tog.config(text="▼  LLM 연결 설정")
        else:
            self._llm_panel.pack_forget()
            self._llm_tog.config(text="▶  LLM 연결 설정")

    def _on_url(self,*_):
        pid=extract_page_id_from_url(self.url_var.get())
        self.pid_lbl.config(text=f"페이지 ID: {pid}" if pid else "페이지 ID: 인식되지 않음",
                            fg=GREEN if pid else RED)

    def _browse(self):
        d=filedialog.askdirectory()
        if d: self.out_var.set(d)

    def _clear_log(self): self.log_box.delete("1.0","end")

    def _save_llm(self):
        k=self.llm_key_var.get().strip(); u=self.llm_url_var.get().strip()
        m=self.llm_model_var.get().strip(); t=self.llm_tok_var.get().strip()
        if not k or not u:
            self._llm_st.config(text="API Key와 Base URL은 필수입니다.",fg=RED); return
        try: save_llm_config(k,u,m,int(t)); self._llm_st.config(text="저장 완료!",fg=GREEN)
        except Exception as e: self._llm_st.config(text=f"실패: {e}",fg=RED)

    def _test_llm(self):
        self._save_llm(); self._llm_st.config(text="연결 테스트 중...",fg=LOG_BLUE)
        def _t():
            try:
                r=_client().chat.completions.create(
                    model=LLM_MODEL,max_tokens=30,temperature=0,
                    messages=[{"role":"user","content":"안녕"}])
                ans=r.choices[0].message.content.strip()[:40]
                self.root.after(0,lambda:self._llm_st.config(text=f"연결 성공! {ans}",fg=GREEN))
            except Exception as e:
                err=str(e); self.root.after(0,lambda:self._llm_st.config(text=f"실패: {err}",fg=RED))
        threading.Thread(target=_t,daemon=True).start()

    def log(self,msg:str):
        try:
            ts=datetime.now().strftime("%H:%M:%S"); msg=msg.strip()
            if not msg: return
            self.log_box.insert("end",f"[{ts}] ","ts")
            if msg.startswith("OK:"):
                self.log_box.insert("end","✓ "+msg[3:].strip()+"\n","ok")
            elif msg.startswith("ERR:"):
                self.log_box.insert("end","✗ "+msg[4:].strip()+"\n","err")
            elif msg.startswith("WARN:"):
                self.log_box.insert("end","⚠ "+msg[5:].strip()+"\n","warn")
            elif msg.startswith("DONE:"):
                self.log_box.insert("end","★ "+msg[5:].strip()+"\n","done")
            else:
                self.log_box.insert("end",(msg[5:] if msg.startswith("INFO:") else msg)+"\n","info")
            self.log_box.see("end"); self.root.update_idletasks()
        except: pass

    def _set_status(self,text,color=TEXT_MUTED):
        self._st_lbl.config(text=f"  {text}",fg=color)
        dot={"#16A34A":"#86EFAC","#DC2626":"#FCA5A5"}.get(color,"#93C5FD")
        self._hdr_status.config(text=f"●  {text}",fg=dot)

    def _run(self):
        pid=extract_page_id_from_url(self.url_var.get())
        if not pid:
            messagebox.showerror("오류","URL에서 페이지 ID를 찾을 수 없습니다."); return
        if self.llm_var.get() and not _llm_ready():
            if not messagebox.askyesno("LLM 미설정","LLM 설정 없음. MD만 수집할까요?"): return
            self.llm_var.set(False)
        self.run_btn.config_state("disabled"); self.progress.start(12)
        self._set_status("수집 중...",BLUE); self._clear_log()
        threading.Thread(target=self._worker,args=(pid,),daemon=True).start()

    def _worker(self,page_id):
        from playwright.sync_api import sync_playwright
        save_dir=os.path.join(self.out_var.get(),f"page_{page_id}")
        os.makedirs(save_dir,exist_ok=True)
        try:
            with sync_playwright() as p:
                browser=p.chromium.launch_persistent_context(
                    user_data_dir=USER_DATA_DIR,headless=False,
                    viewport={"width":1280,"height":720})
                page=browser.new_page()
                page.goto(f"{BASE_URL}/pages/viewpage.action?pageId={page_id}")
                try: page.wait_for_load_state("networkidle",timeout=30000)
                except: pass
                if not messagebox.askokcancel("로그인 확인",
                        "Confluence에 로그인되어 있나요?\n[확인] 진행  [취소] 중단"):
                    browser.close(); return
                page.reload(wait_until="networkidle")
                self.log("="*52)
                self.log(f"INFO: 수집 시작 | 페이지 ID: {page_id}")
                self.log(f"INFO: LLM={self.llm_var.get()} | Vision={self.vision_var.get()} | Word={self.word_var.get()}")
                self.log("="*52)

                process_child_pages_recursive(
                    page,page_id,save_dir,
                    depth=self.depth_var.get(),
                    use_llm=self.llm_var.get(),
                    use_vision=self.vision_var.get(),
                    callback=self.log)

                self.log("INFO: 종합 보고서 생성 중...")
                report_text=generate_report_from_md_files(
                    save_dir,use_llm=self.llm_var.get(),callback=self.log)

                ts=datetime.now().strftime("%Y%m%d_%H%M")

                # MD 저장
                md_path=os.path.join(save_dir,f"{page_id}_종합보고서_{ts}.md")
                with open(md_path,"w",encoding="utf-8") as f: f.write(report_text)
                self.log(f"OK: MD 보고서 저장: {md_path}")

                # Word 저장
                if self.word_var.get():
                    self.log("INFO: Word 문서 변환 중...")
                    docx_path=os.path.join(save_dir,f"{page_id}_종합보고서_{ts}.docx")
                    create_word_report(
                        report_text=report_text,
                        title=f"Confluence 종합 보고서 (페이지 {page_id})",
                        output_path=docx_path,
                        use_charts=self.chart_var.get(),
                        callback=self.log)
                    self.log(f"OK: Word 보고서 저장: {docx_path}")

                browser.close()
                self.log("DONE: 모든 작업 완료!")
                self._set_status("완료",GREEN)
                messagebox.showinfo("완료",f"완료!\n저장 위치: {os.path.abspath(save_dir)}")

        except Exception as e:
            self.log(f"ERR: {e}\n{traceback.format_exc()}")
            self._set_status("오류 발생",RED)
            messagebox.showerror("오류",str(e))
        finally:
            self.root.after(0,lambda:(self.progress.stop(),self.run_btn.config_state("normal")))


def main():
    root=tk.Tk()
    style=ttk.Style()
    try: style.theme_use("clam")
    except: pass
    style.configure("TProgressbar",troughcolor=BORDER,background=BLUE,
                    bordercolor=BORDER,lightcolor=BLUE,darkcolor=BLUE_DK)
    ConfluenceGUI(root); root.mainloop()

if __name__=="__main__":
    main()
