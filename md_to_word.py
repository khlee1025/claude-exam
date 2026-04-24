"""
md_to_word.py - Markdown → Word(.docx) 변환기
LLM 없이 순수 MD 파싱으로 Word 문서 생성
"""
import subprocess, sys

def _install(pkg):
    subprocess.check_call([sys.executable,"-m","pip","install",pkg,"-q"],
                          stdout=subprocess.DEVNULL,stderr=subprocess.DEVNULL)

try: from docx import Document
except ImportError: print("python-docx 설치 중..."); _install("python-docx"); from docx import Document

import os, re, glob, threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ─── 색상 ─────────────────────────────────────
BLUE="#1428A0"; BLUE_LT="#E8EAFF"; BG="#F0F2F7"
CARD="#FFFFFF"; BORDER="#D0D5E0"; GREEN="#16A34A"
RED="#DC2626"; TEXT="#1A1A2E"; TEXT_MUTED="#6B7280"
LOG_BG="#0D1117"; LOG_FG="#C9D1D9"
LOG_GREEN="#3FB950"; LOG_RED="#FF7B72"; LOG_BLUE="#79C0FF"

# ─── MD → Word 변환 엔진 ───────────────────────
class MDConverter:
    HEADING_COLORS = {
        1: RGBColor(0x14,0x28,0xA0),  # 삼성 블루
        2: RGBColor(0x1F,0x49,0x7D),
        3: RGBColor(0x2E,0x74,0xB5),
        4: RGBColor(0x44,0x72,0xC4),
    }
    HEADING_SIZES = {1:18, 2:15, 3:13, 4:12}

    def __init__(self):
        self.doc = None

    def _new_doc(self) -> Document:
        doc = Document()
        for section in doc.sections:
            section.page_width  = Cm(21)
            section.page_height = Cm(29.7)
            section.left_margin = section.right_margin = Cm(2.5)
            section.top_margin  = section.bottom_margin = Cm(2.5)
        # 기본 스타일
        style = doc.styles['Normal']
        style.font.name = "맑은 고딕"
        style.font.size = Pt(11)
        return doc

    def _add_heading(self, doc, text, level):
        level = min(max(level,1),4)
        p = doc.add_heading("", level=level)
        run = p.add_run(text)
        run.font.name = "맑은 고딕"
        run.font.size = Pt(self.HEADING_SIZES[level])
        run.font.color.rgb = self.HEADING_COLORS[level]
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def _add_paragraph(self, doc, text, bold=False, italic=False, size=11):
        if not text.strip(): return doc.add_paragraph()
        p = doc.add_paragraph()
        self._add_inline(p, text, base_bold=bold, base_italic=italic, size=size)
        return p

    def _add_inline(self, para, text, base_bold=False, base_italic=False, size=11):
        """**bold**, *italic*, `code` 인라인 처리"""
        # 패턴: **bold**, *italic*, `code`
        pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
        last = 0
        for m in pattern.finditer(text):
            # 앞 일반 텍스트
            if m.start() > last:
                run = para.add_run(text[last:m.start()])
                run.font.name="맑은 고딕"; run.font.size=Pt(size)
                run.bold=base_bold; run.italic=base_italic
            raw = m.group(0)
            if raw.startswith("**"):
                run = para.add_run(m.group(2))
                run.bold=True; run.italic=base_italic
            elif raw.startswith("*"):
                run = para.add_run(m.group(3))
                run.italic=True; run.bold=base_bold
            elif raw.startswith("`"):
                run = para.add_run(m.group(4))
                run.font.name="Consolas"; run.font.size=Pt(size-1)
                run.font.color.rgb = RGBColor(0xC0,0x39,0x2B)
            run.font.size = Pt(size)
            last = m.end()
        if last < len(text):
            run = para.add_run(text[last:])
            run.font.name="맑은 고딕"; run.font.size=Pt(size)
            run.bold=base_bold; run.italic=base_italic

    def _add_bullet(self, doc, text, level=0):
        style = "List Bullet" if level==0 else "List Bullet 2"
        p = doc.add_paragraph(style=style)
        self._add_inline(p, text.lstrip("-•* ").strip())
        return p

    def _add_numbered(self, doc, text):
        p = doc.add_paragraph(style="List Number")
        clean = re.sub(r"^\d+\.\s*","",text)
        self._add_inline(p, clean)
        return p

    def _add_table(self, doc, headers, rows):
        cols = max(len(headers), max((len(r) for r in rows), default=1))
        table = doc.add_table(rows=1+len(rows), cols=cols)
        table.style = "Table Grid"
        # 헤더
        hrow = table.rows[0]
        for i, h in enumerate(headers[:cols]):
            cell = hrow.cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.bold = True; run.font.name="맑은 고딕"; run.font.size=Pt(10)
            # 헤더 배경색
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = tc.makeelement(qn('w:shd'),{
                qn('w:val'):'clear', qn('w:color'):'auto',
                qn('w:fill'):'1428A0'})
            tcPr.append(shd)
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        # 데이터
        for ri, row in enumerate(rows):
            for ci in range(cols):
                val = row[ci] if ci < len(row) else ""
                cell = table.rows[ri+1].cells[ci]
                cell.text = ""
                run = cell.paragraphs[0].add_run(val)
                run.font.name="맑은 고딕"; run.font.size=Pt(10)
                # 짝수 행 연한 배경
                if ri % 2 == 1:
                    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
                    shd = tc.makeelement(qn('w:shd'),{
                        qn('w:val'):'clear', qn('w:color'):'auto',
                        qn('w:fill'):'F0F4FF'})
                    tcPr.append(shd)
        doc.add_paragraph()
        return table

    def _add_code_block(self, doc, code):
        p = doc.add_paragraph()
        run = p.add_run(code)
        run.font.name="Consolas"; run.font.size=Pt(9)
        run.font.color.rgb = RGBColor(0x20,0x20,0x20)
        pPr = p._p.get_or_add_pPr()
        from docx.oxml import OxmlElement
        pBdr = OxmlElement('w:pBdr')
        for side in ['top','bottom','left','right']:
            bdr = OxmlElement(f'w:{side}')
            bdr.set(qn('w:val'),'single'); bdr.set(qn('w:sz'),'4')
            bdr.set(qn('w:space'),'2'); bdr.set(qn('w:color'),'C0C0C0')
            pBdr.append(bdr)
        pPr.append(pBdr)
        from docx.oxml.ns import qn as _qn
        shd = OxmlElement('w:shd')
        shd.set(_qn('w:val'),'clear'); shd.set(_qn('w:color'),'auto')
        shd.set(_qn('w:fill'),'F6F8FA')
        pPr.append(shd)
        return p

    def _add_hr(self, doc):
        p = doc.add_paragraph()
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6')
        bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'1428A0')
        pBdr.append(bottom); pPr.append(pBdr)

    def convert(self, md_text: str, title: str = "", add_cover: bool = True) -> Document:
        doc = self._new_doc()

        # 커버 페이지
        if add_cover and title:
            # 빈 줄 2개
            doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.font.name="맑은 고딕"; run.font.size=Pt(22); run.bold=True
            run.font.color.rgb = RGBColor(0x14,0x28,0xA0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"작성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            run2.font.name="맑은 고딕"; run2.font.size=Pt(11)
            run2.font.color.rgb = RGBColor(0x6B,0x72,0x80)
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()

        lines = md_text.split("\n")
        i = 0
        in_code = False
        code_lines = []
        table_headers = []
        table_rows = []

        def flush_table():
            nonlocal table_headers, table_rows
            if table_headers or table_rows:
                self._add_table(doc, table_headers, table_rows)
                table_headers = []; table_rows = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # 코드 블록
            if stripped.startswith("```"):
                if not in_code:
                    flush_table()
                    in_code = True; code_lines = []
                else:
                    in_code = False
                    self._add_code_block(doc, "\n".join(code_lines))
                i += 1; continue

            if in_code:
                code_lines.append(line); i += 1; continue

            # 빈 줄
            if not stripped:
                flush_table(); doc.add_paragraph(); i += 1; continue

            # 수평선
            if re.match(r'^[-*_]{3,}$', stripped):
                flush_table(); self._add_hr(doc); i += 1; continue

            # 헤딩
            m = re.match(r'^(#{1,6})\s+(.*)', stripped)
            if m:
                flush_table()
                self._add_heading(doc, m.group(2), len(m.group(1)))
                i += 1; continue

            # MD 테이블
            if stripped.startswith("|"):
                cells = [c.strip() for c in stripped.strip("|").split("|")]
                if all(re.match(r'^[-:]+$', c) for c in cells if c):
                    i += 1; continue  # 구분선 스킵
                if not table_headers:
                    table_headers = cells
                else:
                    table_rows.append(cells)
                i += 1; continue

            # 테이블이 끊기면 flush
            if table_headers and not stripped.startswith("|"):
                flush_table()

            # 불릿 리스트
            if re.match(r'^[-*•]\s+', stripped):
                flush_table()
                self._add_bullet(doc, stripped)
                i += 1; continue

            # 들여쓰기 불릿 (2칸+)
            if re.match(r'^\s{2,}[-*•]\s+', line):
                flush_table()
                self._add_bullet(doc, stripped, level=1)
                i += 1; continue

            # 번호 리스트
            if re.match(r'^\d+\.\s+', stripped):
                flush_table()
                self._add_numbered(doc, stripped)
                i += 1; continue

            # 인용 (> )
            if stripped.startswith(">"):
                flush_table()
                p = doc.add_paragraph()
                run = p.add_run(stripped.lstrip("> "))
                run.font.name="맑은 고딕"; run.font.size=Pt(10)
                run.font.color.rgb = RGBColor(0x6B,0x72,0x80); run.italic=True
                from docx.oxml import OxmlElement
                pPr = p._p.get_or_add_pPr()
                ind = OxmlElement('w:ind'); ind.set(qn('w:left'),'720')
                pPr.append(ind)
                i += 1; continue

            # 일반 텍스트
            flush_table()
            self._add_paragraph(doc, stripped)
            i += 1

        flush_table()

        # 푸터
        for section in doc.sections:
            footer = section.footer
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.text = f"  {title}  |  {datetime.now().strftime('%Y-%m-%d')}  "
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in fp.runs:
                run.font.size=Pt(8); run.font.name="맑은 고딕"
                run.font.color.rgb=RGBColor(0x9C,0xA3,0xAF)

        return doc

    def convert_file(self, md_path: str, out_path: str = None,
                     add_cover: bool = True, callback=None) -> str:
        with open(md_path, encoding="utf-8") as f:
            md_text = f.read()
        title = os.path.basename(md_path).replace(".md","")
        if callback: callback(f"INFO: 변환 중: {title}")
        doc = self.convert(md_text, title=title, add_cover=add_cover)
        if out_path is None:
            out_path = md_path.replace(".md",".docx")
        doc.save(out_path)
        if callback: callback(f"OK: 저장 완료 → {out_path}")
        return out_path

    def convert_folder(self, folder: str, out_folder: str = None,
                       merge: bool = False, add_cover: bool = True,
                       callback=None) -> list:
        md_files = sorted(glob.glob(os.path.join(folder,"**","*.md"), recursive=True))
        md_files = [f for f in md_files
                    if "보고서" not in os.path.basename(f) or not merge]
        if not md_files:
            if callback: callback("WARN: MD 파일 없음"); return []
        if callback: callback(f"INFO: MD 파일 {len(md_files)}개 발견")

        if merge:
            # 하나로 합치기
            combined = ""
            for fp in md_files:
                fname = os.path.basename(fp).replace(".md","")
                with open(fp,encoding="utf-8") as f: text=f.read()
                combined += f"\n\n---\n\n# {fname}\n\n{text}"
            title = os.path.basename(folder) + " 통합 보고서"
            doc = self.convert(combined, title=title, add_cover=add_cover)
            out = os.path.join(out_folder or folder, f"{title}.docx")
            doc.save(out)
            if callback: callback(f"OK: 통합 문서 저장 → {out}")
            return [out]
        else:
            results = []
            odir = out_folder or folder
            os.makedirs(odir, exist_ok=True)
            for fp in md_files:
                fname = os.path.basename(fp).replace(".md","")
                out = os.path.join(odir, fname+".docx")
                try:
                    self.convert_file(fp, out, add_cover=add_cover, callback=callback)
                    results.append(out)
                except Exception as e:
                    if callback: callback(f"ERR: {fname} 실패 - {e}")
            return results


# ─── GUI ──────────────────────────────────────
class MDToWordGUI:
    def __init__(self, root):
        self.root=root
        self.root.title("MD → Word 변환기")
        self.root.geometry("700x580"); self.root.configure(bg=BG)

        self.mode_var  = tk.StringVar(value="files")   # files | folder
        self.merge_var = tk.BooleanVar(value=False)
        self.cover_var = tk.BooleanVar(value=True)
        self.out_var   = tk.StringVar()
        self._files    = []

        self._build()

    def _build(self):
        # 헤더
        hdr=tk.Frame(self.root,bg=BLUE,height=52); hdr.pack(fill="x"); hdr.pack_propagate(False)
        tk.Label(hdr,text="  MD → Word 변환기",bg=BLUE,fg="white",
                 font=("Malgun Gothic",13,"bold")).pack(side="left",padx=16)
        tk.Label(hdr,text=" LLM 없음 ",bg="#1E3A8A",fg="#93C5FD",
                 font=("Malgun Gothic",8),padx=6).pack(side="left",pady=16)

        # 메인 카드
        card=tk.Frame(self.root,bg=CARD,highlightthickness=1,highlightbackground=BORDER)
        card.pack(fill="x",padx=14,pady=10)
        f=tk.Frame(card,bg=CARD); f.pack(fill="x",padx=16,pady=12)

        # 모드 선택
        self._lbl(f,"변환 방식")
        mf=tk.Frame(f,bg=CARD); mf.pack(anchor="w",pady=(0,8))
        tk.Radiobutton(mf,text="파일 직접 선택",variable=self.mode_var,value="files",
                       bg=CARD,fg=TEXT,activebackground=CARD,font=("Malgun Gothic",9),
                       command=self._on_mode).pack(side="left",padx=(0,16))
        tk.Radiobutton(mf,text="폴더 전체",variable=self.mode_var,value="folder",
                       bg=CARD,fg=TEXT,activebackground=CARD,font=("Malgun Gothic",9),
                       command=self._on_mode).pack(side="left")

        # 파일/폴더 선택
        self._lbl(f,"MD 파일 / 폴더")
        sf=tk.Frame(f,bg=CARD); sf.pack(fill="x",pady=(0,4))
        self.path_lbl=tk.Label(sf,text="(선택 없음)",bg=BLUE_LT,fg=BLUE,
                               font=("Malgun Gothic",9),anchor="w",padx=8,pady=4)
        self.path_lbl.pack(side="left",fill="x",expand=True)
        self.pick_btn=tk.Button(sf,text="파일 선택",command=self._pick_files,
                                bg=BLUE,fg="white",relief="flat",cursor="hand2",
                                font=("Malgun Gothic",9),padx=10,pady=4)
        self.pick_btn.pack(side="left",padx=(6,0))

        # 출력 폴더
        self._lbl(f,"출력 폴더 (비워두면 원본 위치에 저장)")
        of=tk.Frame(f,bg=CARD); of.pack(fill="x",pady=(0,8))
        tk.Entry(of,textvariable=self.out_var,font=("Malgun Gothic",9),
                 bd=1,relief="solid",bg="white",fg=TEXT).pack(side="left",fill="x",expand=True)
        tk.Button(of,text="찾기",command=self._pick_out,bg=BG,fg=TEXT,
                  relief="flat",cursor="hand2",font=("Malgun Gothic",9),padx=6).pack(side="left",padx=(4,0))

        # 옵션
        tk.Frame(f,bg=BORDER,height=1).pack(fill="x",pady=6)
        optf=tk.Frame(f,bg=CARD); optf.pack(fill="x")
        self.merge_chk=tk.Checkbutton(optf,text="폴더 내 MD 파일 하나로 합치기",
                       variable=self.merge_var,bg=CARD,fg=TEXT,
                       activebackground=CARD,font=("Malgun Gothic",9),
                       selectcolor=BLUE_LT,state="disabled")
        self.merge_chk.pack(side="left",padx=(0,16))
        tk.Checkbutton(optf,text="표지 페이지 추가",variable=self.cover_var,
                       bg=CARD,fg=TEXT,activebackground=CARD,
                       font=("Malgun Gothic",9),selectcolor=BLUE_LT).pack(side="left")

        # 변환 버튼
        tk.Frame(card,bg=BORDER,height=1).pack(fill="x",padx=16)
        bf=tk.Frame(card,bg=CARD); bf.pack(pady=12,padx=16,fill="x")
        self.convert_btn=tk.Button(bf,text="▶  Word로 변환",command=self._run,
                                   bg=BLUE,fg="white",relief="flat",cursor="hand2",
                                   font=("Malgun Gothic",10,"bold"),padx=20,pady=8)
        self.convert_btn.pack(side="left")
        self.progress=ttk.Progressbar(bf,mode="indeterminate",length=300)
        self.progress.pack(side="left",padx=16)

        # 로그
        lc=tk.Frame(self.root,bg=LOG_BG,highlightthickness=1,highlightbackground="#30363D")
        lc.pack(fill="both",expand=True,padx=14,pady=(0,10))
        lh=tk.Frame(lc,bg="#161B22"); lh.pack(fill="x")
        tk.Label(lh,text="  로그",bg="#161B22",fg=LOG_BLUE,
                 font=("Consolas",9,"bold")).pack(side="left",pady=5,padx=8)
        tk.Button(lh,text="지우기",command=lambda:self.log_box.delete("1.0","end"),
                  bg="#161B22",fg="#6E7681",relief="flat",cursor="hand2",
                  font=("Consolas",8),padx=6).pack(side="right",padx=8)
        tf=tk.Frame(lc,bg=LOG_BG); tf.pack(fill="both",expand=True,padx=4,pady=4)
        sb=tk.Scrollbar(tf,bg=LOG_BG,troughcolor=LOG_BG); sb.pack(side="right",fill="y")
        self.log_box=tk.Text(tf,bg=LOG_BG,fg=LOG_FG,font=("Consolas",9),
                             wrap="word",yscrollcommand=sb.set,bd=0,relief="flat",
                             selectbackground="#264F78")
        self.log_box.pack(fill="both",expand=True); sb.config(command=self.log_box.yview)
        self.log_box.tag_config("ok",foreground=LOG_GREEN)
        self.log_box.tag_config("err",foreground=LOG_RED)
        self.log_box.tag_config("info",foreground=LOG_BLUE)
        self.log_box.tag_config("ts",foreground="#484F58")
        self.log_box.tag_config("done",foreground=LOG_GREEN,font=("Consolas",9,"bold"))

        # 상태바
        bar=tk.Frame(self.root,bg="#E5E7EB",height=22); bar.pack(fill="x",side="bottom")
        bar.pack_propagate(False)
        self._st=tk.Label(bar,text="  준비됨",bg="#E5E7EB",fg=TEXT_MUTED,
                          font=("Malgun Gothic",8),anchor="w")
        self._st.pack(side="left",fill="both",expand=True)

    def _lbl(self,p,t):
        tk.Label(p,text=t,bg=CARD,fg=TEXT_MUTED,font=("Malgun Gothic",8)).pack(anchor="w",pady=(4,1))

    def _on_mode(self):
        mode=self.mode_var.get()
        self.pick_btn.config(text="파일 선택" if mode=="files" else "폴더 선택")
        self.merge_chk.config(state="normal" if mode=="folder" else "disabled")
        self._files=[]; self.path_lbl.config(text="(선택 없음)")

    def _pick_files(self):
        if self.mode_var.get()=="files":
            paths=filedialog.askopenfilenames(filetypes=[("MD 파일","*.md"),("모든 파일","*.*")])
            if paths:
                self._files=list(paths)
                self.path_lbl.config(text=f"{len(paths)}개 파일 선택됨")
        else:
            d=filedialog.askdirectory()
            if d:
                self._files=[d]
                self.path_lbl.config(text=d)

    def _pick_out(self):
        d=filedialog.askdirectory()
        if d: self.out_var.set(d)

    def log(self,msg:str):
        try:
            from datetime import datetime as dt
            ts=dt.now().strftime("%H:%M:%S"); msg=msg.strip()
            if not msg: return
            self.log_box.insert("end",f"[{ts}] ","ts")
            if msg.startswith("OK:"):
                self.log_box.insert("end","✓ "+msg[3:].strip()+"\n","ok")
            elif msg.startswith("ERR:"):
                self.log_box.insert("end","✗ "+msg[4:].strip()+"\n","err")
            elif msg.startswith("DONE:"):
                self.log_box.insert("end","★ "+msg[5:].strip()+"\n","done")
            else:
                self.log_box.insert("end",(msg[5:] if msg.startswith("INFO:") else msg)+"\n","info")
            self.log_box.see("end"); self.root.update_idletasks()
        except: pass

    def _run(self):
        if not self._files:
            messagebox.showerror("오류","파일 또는 폴더를 선택하세요."); return
        self.convert_btn.config(state="disabled"); self.progress.start(12)
        self._st.config(text="  변환 중...",fg=BLUE)
        threading.Thread(target=self._worker,daemon=True).start()

    def _worker(self):
        conv=MDConverter()
        out=self.out_var.get().strip() or None
        results=[]
        try:
            if self.mode_var.get()=="files":
                for fp in self._files:
                    try:
                        op=os.path.join(out,os.path.basename(fp).replace(".md",".docx")) if out else None
                        res=conv.convert_file(fp,op,add_cover=self.cover_var.get(),callback=self.log)
                        results.append(res)
                    except Exception as e:
                        self.log(f"ERR: {os.path.basename(fp)} - {e}")
            else:
                folder=self._files[0]
                results=conv.convert_folder(folder,out,
                    merge=self.merge_var.get(),
                    add_cover=self.cover_var.get(),callback=self.log)

            self.log(f"DONE: 완료! {len(results)}개 Word 파일 생성")
            self._st.config(text=f"  완료 ({len(results)}개 생성)",fg=GREEN)
            messagebox.showinfo("완료",f"Word 파일 {len(results)}개 생성 완료!")
        except Exception as e:
            self.log(f"ERR: {e}")
            self._st.config(text="  오류 발생",fg=RED)
        finally:
            self.root.after(0,lambda:(self.progress.stop(),
                                      self.convert_btn.config(state="normal")))


def main():
    root=tk.Tk()
    style=ttk.Style()
    try: style.theme_use("clam")
    except: pass
    MDToWordGUI(root); root.mainloop()

if __name__=="__main__":
    main()
